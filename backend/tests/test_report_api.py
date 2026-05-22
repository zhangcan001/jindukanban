from datetime import date
import json
from urllib.parse import unquote

from fastapi.testclient import TestClient
from docx import Document

from app.database import SessionLocal
from app.main import app
from app.models.import_validation_issue import ImportValidationIssue
from app.models.import_batch import ImportBatch
from app.models.baseline_plan import BaselinePlan
from app.models.progress_item import ProgressItem
from app.models.project import Project
from app.models.rectification_item import RectificationItem
from app.models.report_export_record import ReportExportRecord
from app.models.warning_record import WarningRecord
from app.services.report_service import REPORT_TYPES
from openpyxl import load_workbook
from io import BytesIO


def test_report_type_registry_contains_unified_export_types() -> None:
    assert REPORT_TYPES["dashboard_excel"]["label"] == "当前看板 Excel"
    assert REPORT_TYPES["weekly_word"]["label"] == "Word 周报"
    assert REPORT_TYPES["weekly_pdf"]["label"] == "PDF 周报"
    assert REPORT_TYPES["weekly_pdf"]["extension"] == "pdf"
    assert REPORT_TYPES["delay_rectification_excel"]["label"] == "滞后项整改清单"
    assert REPORT_TYPES["delay_rectification_excel"]["extension"] == "xlsx"


def test_reports_exclude_not_started_by_plan_from_delayed_lists() -> None:
    db = SessionLocal()
    try:
        project = Project(name="报表计划开始过滤", project_type="测试")
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="report.csv", status="published", data_date=date(2026, 5, 18))
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    task_name="未到计划开始",
                    actual_percent=0,
                    planned_percent=40,
                    progress_deviation=-40,
                    planned_start_date=date(2026, 5, 20),
                    planned_finish_date=date(2026, 5, 30),
                    status="not_started_by_plan",
                ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    task_name="已到计划开始",
                    actual_percent=0,
                    planned_percent=40,
                    progress_deviation=-40,
                    planned_start_date=date(2026, 5, 8),
                    planned_finish_date=date(2026, 5, 28),
                    status="seriously_delayed",
                ),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        delayed_report = client.get(f"/api/projects/{project_id}/reports/delayed-ranking", params={"batch_id": batch_id})
        dashboard_report = client.get(f"/api/projects/{project_id}/reports/dashboard_excel", params={"batch_id": batch_id})

    assert delayed_report.status_code == 200
    workbook = load_workbook(BytesIO(delayed_report.content), read_only=True)
    delayed_text = "\n".join(str(cell.value) for row in workbook.active.iter_rows() for cell in row if cell.value is not None)
    assert "已到计划开始" in delayed_text
    assert "未到计划开始" not in delayed_text

    assert dashboard_report.status_code == 200
    dashboard_workbook = load_workbook(BytesIO(dashboard_report.content), read_only=True)
    dashboard_text = "\n".join(
        str(cell.value)
        for row in dashboard_workbook["滞后项清单"].iter_rows()
        for cell in row
        if cell.value is not None
    )
    assert "已到计划开始" in dashboard_text
    assert "未到计划开始" not in dashboard_text


def test_report_exports_create_xlsx_files_and_records() -> None:
    db = SessionLocal()
    try:
        project = Project(name="报表导出测试", project_type="测试")
        db.add(project)
        db.flush()
        baseline = BaselinePlan(project_id=project.id, name="当前计划", is_default=True)
        db.add(baseline)
        db.flush()
        batch = ImportBatch(
            project_id=project.id,
            file_name="report.csv",
            status="published",
            data_date=date(2026, 5, 13),
            imported_count=2,
            data_quality_score=88,
            baseline_plan_id=baseline.id,
        )
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="桥架安装",
                    wbs_code="JD.01",
                    discipline="电气",
                    unit="米",
                    total_quantity=100,
                    actual_quantity=60,
                    planned_quantity=80,
                    actual_percent=60,
                    planned_percent=80,
                    progress_deviation=-20,
                    status="seriously_delayed",
                ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="给水管安装",
                    wbs_code="JD.02",
                    discipline="给排水",
                    unit="米",
                    total_quantity=120,
                    actual_quantity=90,
                    planned_quantity=90,
                    actual_percent=75,
                    planned_percent=75,
                    progress_deviation=0,
                    status="normal",
                ),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
        baseline_id = baseline.id
    finally:
        db.close()

    with TestClient(app) as client:
        for report_type in ("overview", "delayed-ranking", "discipline-summary", "progress-items"):
            response = client.get(
                f"/api/projects/{project_id}/reports/{report_type}",
                params={"batch_id": batch_id, "baseline_plan_id": baseline_id},
            )
            assert response.status_code == 200
            assert response.content[:2] == b"PK"
            assert "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" in response.headers["content-type"]
            assert "attachment;" in response.headers["content-disposition"]
            workbook = load_workbook(BytesIO(response.content), read_only=True)
            try:
                values = [cell.value for row in workbook.active.iter_rows(min_row=1, max_row=7) for cell in row]
                assert "批次绑定计划基线" in values
                assert "当前查看计划基线" in values
                assert "是否与批次绑定基线一致" in values
                assert "当前计划" in values
            finally:
                workbook.close()

        exports_response = client.get(f"/api/projects/{project_id}/reports/exports")
        assert exports_response.status_code == 200
        assert len(exports_response.json()) == 4

    db = SessionLocal()
    try:
        records = db.query(ReportExportRecord).filter(ReportExportRecord.project_id == project_id).all()
        assert len(records) == 4
        assert all(record.file_name and record.file_path and record.batch_id == batch_id for record in records)
    finally:
        db.close()


def test_dashboard_export_returns_xlsx_with_dashboard_sheets() -> None:
    db = SessionLocal()
    try:
        project = Project(name="汽车科技总部产业园", project_type="测试")
        db.add(project)
        db.flush()
        baseline = BaselinePlan(project_id=project.id, name="当前计划", is_default=True)
        db.add(baseline)
        db.flush()
        batch = ImportBatch(
            project_id=project.id,
            file_name="进度导入模板-汽车科技总部产业园-2026-05-12.xlsx",
            sheet_name="机电单位",
            status="published",
            data_date=date(2026, 5, 12),
            imported_count=4,
            warning_count=2,
            error_count=1,
            data_quality_score=65,
            baseline_plan_id=baseline.id,
        )
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="桥架安装",
                    building="A1",
                    floor="1层",
                    discipline="机电",
                    system_name="强电",
                    unit="米",
                    actual_percent=50,
                    planned_percent=70,
                    progress_deviation=-20,
                    status="seriously_delayed",
                    planned_start_date=date(2026, 5, 1),
                    planned_finish_date=date(2026, 5, 15),
                    weight=0.25,
                    extra_fields='{"权重": "25%"}',
                ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="风管安装",
                    building="A1",
                    floor="10层",
                    discipline="机电",
                    system_name="通风",
                    unit="个",
                    actual_percent=80,
                        planned_percent=75,
                        progress_deviation=5,
                        status="ahead",
                        planned_start_date=date(2026, 5, 1),
                        planned_finish_date=date(2026, 5, 16),
                    ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="给水管安装",
                    building="A2",
                    floor="2层",
                    discipline="给排水",
                    system_name="给水",
                    unit="米",
                    actual_percent=90,
                        planned_percent=90,
                        progress_deviation=0,
                        status="normal",
                        planned_start_date=date(2026, 5, 1),
                        planned_finish_date=date(2026, 5, 13),
                    ),
            ]
        )
        db.add(ImportValidationIssue(batch_id=batch.id, level="warning", code="SUMMARY_ROW_SKIPPED", message="skip summary"))
        db.add(ImportValidationIssue(batch_id=batch.id, level="error", code="percent_out_of_range", message="bad percent"))
        db.add(WarningRecord(project_id=project.id, batch_id=batch.id, level="warning", title="测试预警", message="需要处理"))
        db.add(WarningRecord(project_id=project.id, batch_id=batch.id, level="warning", title="数据质量评分偏低", message="当前数据质量评分 65.0，低于阈值 70.0"))
        db.commit()
        project_id = project.id
        batch_id = batch.id
        baseline_id = baseline.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(
            f"/api/projects/{project_id}/reports/dashboard-export",
            params={"batch_id": batch_id, "baseline_plan_id": baseline_id, "building": "A1"},
        )

    assert response.status_code == 200
    assert response.content[:2] == b"PK"
    assert "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" in response.headers["content-type"]
    assert "汽车科技总部产业园_进度看板_2026-05-12.xlsx" in unquote(response.headers["content-disposition"])

    workbook = load_workbook(BytesIO(response.content), read_only=True)
    try:
        expected_sheets = [
            "看板总览",
            "专业进度统计",
            "楼层进度统计",
            "楼栋楼层统计",
            "滞后项清单",
            "数据质量与校验问题汇总",
            "整改闭环摘要",
            "整改项明细",
            "进度分析说明",
            "专业进度对比",
            "楼层专业矩阵",
            "楼栋专业矩阵",
            "滞后分布统计",
        ]
        assert workbook.sheetnames == expected_sheets
        overview_values = [cell for row in workbook["看板总览"].iter_rows(values_only=True) for cell in row]
        assert "项目名称" in overview_values
        assert "汽车科技总部产业园" in overview_values
        assert "批次绑定计划基线" in overview_values
        assert "当前查看计划基线" in overview_values
        assert "是否与批次绑定基线一致" in overview_values
        assert "当前计划" in overview_values
        assert "统计口径" in overview_values
        assert "权重统计" in overview_values
        assert "权重来源" in overview_values
        assert "Excel 字段：权重" in overview_values
        assert "当前范围权重合计" in overview_values
        assert "是否归一化" in overview_values

        discipline_rows = list(workbook["专业进度统计"].iter_rows(values_only=True))
        assert any(row[0] == "机电" and "当前分组包含多种单位" in (row[7] or "") for row in discipline_rows)

        floor_rows = list(workbook["楼层进度统计"].iter_rows(values_only=True))
        floors = [row[0] for row in floor_rows[2:] if row[0]]
        assert floors.index("1层") < floors.index("10层")
        assert any(row[0] == "1层" and row[2] == "50.0%" for row in floor_rows)

        building_floor_rows = list(workbook["楼栋楼层统计"].iter_rows(values_only=True))
        assert any(row[0] == "A1" and row[1] == "1层" and row[3] == "50.0%" for row in building_floor_rows)
        assert not any(row[0] == "A2" for row in building_floor_rows[2:])

        delayed_rows = list(workbook["滞后项清单"].iter_rows(values_only=True))
        assert any(
            row[0] == "机电"
            and row[1] == "A1"
            and row[2] == "1层"
            and row[4] == "桥架安装"
            and "【机电】A1 1层 桥架安装" in row[9]
            for row in delayed_rows
        )

        quality_values = [cell for row in workbook["数据质量与校验问题汇总"].iter_rows(values_only=True) for cell in row]
        assert "SUMMARY_ROW_SKIPPED" in quality_values
        assert "percent_out_of_range" in quality_values
        assert any("当前数据质量评分 65.0 分，低于建议阈值 70.0 分" in str(value) for value in quality_values if value)
        delayed_values = [cell for row in workbook["滞后项清单"].iter_rows(values_only=True) for cell in row]
        assert not any("当前数据质量评分 65.0" in str(value) for value in delayed_values if value)

        insight_values = [cell for row in workbook["进度分析说明"].iter_rows(values_only=True) for cell in row]
        assert "总体进度说明" in insight_values
        assert "主要滞后施工项" in insight_values
        assert any(value and "截至 2026-05-12" in str(value) for value in insight_values)
    finally:
        workbook.close()

    db = SessionLocal()
    try:
        record = db.query(ReportExportRecord).filter(ReportExportRecord.project_id == project_id).one()
        assert record.report_type == "dashboard_excel"
        assert record.file_name
        assert record.file_path
        assert record.batch_id == batch_id
    finally:
        db.close()


def test_dashboard_export_returns_project_not_found_code() -> None:
    with TestClient(app) as client:
        response = client.get("/api/projects/999999/reports/dashboard-export")

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "PROJECT_NOT_FOUND"


def test_dashboard_export_returns_clear_error_when_no_published_batch() -> None:
    db = SessionLocal()
    try:
        project = Project(name="无批次项目")
        db.add(project)
        db.commit()
        project_id = project.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(f"/api/projects/{project_id}/reports/dashboard-export")

    assert response.status_code == 404
    assert response.json()["detail"] == {"code": "NO_PUBLISHED_BATCH", "message": "当前暂无可导出数据。"}


def test_weekly_word_returns_docx_with_required_sections_and_delayed_fields() -> None:
    db = SessionLocal()
    try:
        project = Project(name="汽车科技总部产业园", project_type="测试")
        db.add(project)
        db.flush()
        baseline = BaselinePlan(project_id=project.id, name="当前计划", is_default=True)
        db.add(baseline)
        db.flush()
        batch = ImportBatch(
            project_id=project.id,
            file_name="进度导入模板-汽车科技总部产业园-2026-05-12.xlsx",
            sheet_name="消防单位",
            status="published",
            data_date=date(2026, 5, 12),
            imported_count=3,
            warning_count=1,
            error_count=0,
            data_quality_score=65,
            baseline_plan_id=baseline.id,
        )
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="喷淋系统",
                    building="A1",
                    floor="3层",
                    discipline="消防",
                    system_name="喷淋系统",
                    unit="项",
                    actual_percent=58,
                    planned_percent=69,
                        progress_deviation=-11,
                        status="seriously_delayed",
                        planned_start_date=date(2026, 5, 1),
                        planned_finish_date=date(2026, 5, 15),
                        extra_fields=json.dumps({"责任人": "张工"}, ensure_ascii=False),
                    ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="桥架安装",
                    building="A1",
                    floor="10层",
                    discipline="机电",
                    system_name="强电",
                    unit="米",
                    actual_percent=70,
                        planned_percent=75,
                        progress_deviation=-5,
                        status="delayed",
                        planned_start_date=date(2026, 5, 1),
                        planned_finish_date=date(2026, 5, 16),
                    ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="风管安装",
                    building="A2",
                    floor="1层",
                    discipline="暖通",
                    system_name="通风",
                    unit="米",
                    actual_percent=90,
                        planned_percent=80,
                        progress_deviation=10,
                        status="ahead",
                        planned_start_date=date(2026, 5, 1),
                        planned_finish_date=date(2026, 5, 21),
                    ),
            ]
        )
        db.add(ImportValidationIssue(batch_id=batch.id, level="warning", code="INVALID_DATE", message="bad date"))
        db.add(WarningRecord(project_id=project.id, batch_id=batch.id, level="warning", title="测试预警", message="需要处理"))
        db.add(WarningRecord(project_id=project.id, batch_id=batch.id, level="warning", title="数据质量评分偏低", message="当前数据质量评分 65.0，低于阈值 70.0"))
        db.commit()
        project_id = project.id
        batch_id = batch.id
        baseline_id = baseline.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(
            f"/api/projects/{project_id}/reports/weekly-word",
            params={"batch_id": batch_id, "baseline_plan_id": baseline_id},
        )

    assert response.status_code == 200
    assert response.content[:2] == b"PK"
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers["content-type"]
    assert "汽车科技总部产业园_进度周报_2026-05-12.docx" in unquote(response.headers["content-disposition"])

    document = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "工程进度周报" in text
    assert "一、总体进度概况" in text
    assert "二、分专业进度情况" in text
    assert "三、楼层进度情况" in text
    assert "四、楼栋楼层进度情况" in text
    assert "五、主要滞后项" in text
    assert "六、数据质量与校验问题" in text
    assert "七、进度分析说明" in text
    assert "项目进度管理系统自动生成" in text
    assert "当前滞后较明显的专业主要包括" in text
    assert "总体进度说明" in text
    assert "主要滞后施工项" in text
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "统计口径" in table_text
    assert "截至 2026-05-12" in text
    assert "本期关注事项" in text
    assert "建议措施" in text
    assert any("导出时间：" in paragraph.text for section in document.sections for paragraph in section.footer.paragraphs)

    table_texts = ["|".join(cell.text for row in table.rows for cell in row.cells) for table in document.tables]
    all_table_text = "\n".join(table_texts)
    assert "项目名称" in all_table_text
    assert "重点指标表" in text
    assert "滞后项数量" in all_table_text
    assert "严重滞后" in all_table_text
    assert "专业" in all_table_text
    assert "楼栋" in all_table_text
    assert "楼层" in all_table_text
    assert "系统" in all_table_text
    assert "施工项" in all_table_text
    assert "滞后说明" in all_table_text
    assert "消防" in all_table_text
    assert "A1" in all_table_text
    assert "3层" in all_table_text
    assert "喷淋系统" in all_table_text
    assert "【消防】A1 3层 喷淋系统" in all_table_text
    full_doc_text = f"{text}\n{all_table_text}"
    assert "当前数据质量评分 65.0 分，低于建议阈值 70.0 分" in full_doc_text
    assert "当前数据质量评分 65.0，低于阈值 70.0" not in full_doc_text


def test_report_config_defaults_and_weekly_word_advanced_sections_toggle() -> None:
    db = SessionLocal()
    try:
        project = Project(name="report plus alpha", report_config=None)
        db.add(project)
        db.flush()
        batch = ImportBatch(
            project_id=project.id,
            file_name="report-plus.xlsx",
            sheet_name="消防单位",
            status="published",
            data_date=date(2026, 5, 17),
            imported_count=4,
        )
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="喷淋A1", building="A1", floor="1层", discipline="消防", actual_percent=30, planned_percent=60, progress_deviation=-30, status="seriously_delayed", planned_start_date=date(2026, 5, 7), planned_finish_date=date(2026, 5, 17)),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="喷淋A2", building="A2", floor="2层", discipline="消防", actual_percent=60, planned_percent=66, progress_deviation=-6, status="delayed", planned_start_date=date(2026, 5, 4), planned_finish_date=date(2026, 5, 24)),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="桥架A1", building="A1", floor="1层", discipline="机电", actual_percent=81, planned_percent=84, progress_deviation=-3, status="slightly_delayed", planned_start_date=date(2026, 5, 1), planned_finish_date=date(2026, 5, 20)),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="风管A1", building="A1", floor="3层", discipline="暖通", actual_percent=90, planned_percent=80, progress_deviation=10, status="ahead", planned_start_date=date(2026, 5, 1), planned_finish_date=date(2026, 5, 21)),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        defaults = client.get(f"/api/projects/{project_id}/reports/config")
        response = client.get(f"/api/projects/{project_id}/reports/weekly-word", params={"batch_id": batch_id})
        dashboard = client.get(f"/api/projects/{project_id}/reports/dashboard-export", params={"batch_id": batch_id})

        disabled = defaults.json() | {"include_advanced_chart_analysis": False}
        update = client.put(f"/api/projects/{project_id}/reports/config", json=disabled)
        disabled_response = client.get(f"/api/projects/{project_id}/reports/weekly-word", params={"batch_id": batch_id})

    assert defaults.status_code == 200
    assert defaults.json()["include_advanced_chart_analysis"] is True
    assert defaults.json()["weekly_delayed_item_limit"] == 30
    assert defaults.json()["weekly_matrix_summary_limit"] == 10
    assert defaults.json()["show_data_quality_section"] is True
    assert defaults.json()["show_rectification_summary"] is True
    assert defaults.json()["default_export_format"] == "xlsx"
    assert defaults.json()["file_name_include_project_name"] is True
    assert defaults.json()["file_name_include_data_date"] is True

    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join("|".join(cell.text for row in table.rows for cell in row.cells) for table in document.tables)
    assert "八、进阶图表分析" in text
    assert "专业进度对比摘要" in text
    assert "楼层专业矩阵摘要" in text
    assert "楼栋专业矩阵摘要" in text
    assert "滞后分布统计" in text
    assert "专业|任务数|实际进度|应完成进度|进度偏差|滞后数量" in table_text
    assert "楼层|专业|任务数|实际进度|应完成进度|偏差|滞后数量" in table_text
    assert "楼栋|专业|任务数|实际进度|应完成进度|偏差|滞后数量" in table_text
    assert "严重滞后" in table_text
    assert "明显滞后" in table_text
    assert "轻微滞后" in table_text
    assert "正常" in table_text
    assert "超前" in table_text
    assert "None" not in table_text
    assert "null" not in table_text
    assert "undefined" not in table_text

    assert update.status_code == 200
    assert update.json()["include_advanced_chart_analysis"] is False
    disabled_document = Document(BytesIO(disabled_response.content))
    disabled_text = "\n".join(paragraph.text for paragraph in disabled_document.paragraphs)
    assert "八、进阶图表分析" not in disabled_text

    assert dashboard.status_code == 200
    workbook = load_workbook(BytesIO(dashboard.content), read_only=True)
    try:
        assert {"专业进度对比", "楼层专业矩阵", "楼栋专业矩阵", "滞后分布统计"}.issubset(set(workbook.sheetnames))
    finally:
        workbook.close()


def test_report_preview_history_filters_and_structured_errors_for_beta() -> None:
    db = SessionLocal()
    try:
        project = Project(name="beta report center")
        empty_project = Project(name="beta no batch")
        db.add_all([project, empty_project])
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="beta.xlsx", sheet_name="机电单位", status="published", data_date=date(2026, 5, 18))
        db.add(batch)
        db.flush()
        db.add(ProgressItem(project_id=project.id, batch_id=batch.id, task_name="桥架", building="A1", floor="1层", discipline="机电", actual_percent=40, planned_percent=60, progress_deviation=-20))
        db.add(RectificationItem(project_id=project.id, batch_id=batch.id, source_type="manual", task_name="桥架", discipline="机电", building="A1", floor="1层", status="open"))
        db.add(ReportExportRecord(project_id=project.id, batch_id=batch.id, report_type="weekly_word", file_name="beta_weekly.docx", file_path="reports/beta_weekly.docx", data_date=date(2026, 5, 18)))
        db.add(ReportExportRecord(project_id=project.id, batch_id=batch.id, report_type="dashboard_excel", file_name="beta_dashboard.xlsx", file_path="reports/beta_dashboard.xlsx", data_date=date(2026, 5, 18)))
        db.commit()
        project_id = project.id
        empty_project_id = empty_project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        word_preview = client.get(f"/api/projects/{project_id}/reports/preview/weekly_word", params={"batch_id": batch_id})
        dashboard_preview = client.get(f"/api/projects/{project_id}/reports/preview/dashboard_excel", params={"batch_id": batch_id, "building": "A1"})
        rect_preview = client.get(f"/api/projects/{project_id}/reports/preview/rectification_tracking", params={"batch_id": batch_id})
        history = client.get(f"/api/projects/{project_id}/reports/exports", params={"report_type": "weekly_word", "keyword": "weekly", "project_name": "beta"})
        no_batch = client.get(f"/api/projects/{empty_project_id}/reports/weekly-word")
        unknown = client.get(f"/api/projects/{project_id}/reports/preview/unknown_type")
        no_rectifications = client.get(f"/api/projects/{project_id}/rectifications/export", params={"batch_id": batch_id, "discipline": "不存在"})

    assert word_preview.status_code == 200
    word_items = {item["label"]: item["value"] for item in word_preview.json()["items"]}
    assert word_items["项目名称"] == "beta report center"
    assert word_items["是否包含进阶图表分析"] is True
    assert word_items["主要滞后项最大条数"] == 30
    assert word_items["矩阵摘要最大条数"] == 10

    assert dashboard_preview.status_code == 200
    dashboard_items = {item["label"]: item["value"] for item in dashboard_preview.json()["items"]}
    assert "专业进度对比" in dashboard_items["包含 Sheet 列表"]
    assert "楼栋=A1" in dashboard_items["当前筛选条件"]

    assert rect_preview.status_code == 200
    rect_items = {item["label"]: item["value"] for item in rect_preview.json()["items"]}
    assert rect_items["预计导出整改项数量"] == 1

    assert history.status_code == 200
    assert len(history.json()) == 1
    assert history.json()[0]["report_type"] == "weekly_word"

    assert no_batch.status_code == 404
    assert no_batch.json()["detail"] == {"code": "NO_PUBLISHED_BATCH", "message": "当前暂无可导出数据。"}
    assert unknown.status_code == 404
    assert unknown.json()["detail"] == {"code": "REPORT_TYPE_NOT_FOUND", "message": "报表类型不存在或未注册。"}
    assert no_rectifications.status_code == 404
    assert no_rectifications.json()["detail"] == {"code": "NO_RECTIFICATIONS_FOR_FILTER", "message": "当前筛选条件下暂无整改项。"}


def test_delay_rectification_export_returns_xlsx_with_tracking_fields_and_sorted_delays() -> None:
    db = SessionLocal()
    try:
        project = Project(name="整改清单项目")
        db.add(project)
        db.flush()
        baseline = BaselinePlan(project_id=project.id, name="消防专项计划", is_default=True)
        db.add(baseline)
        db.flush()
        batch = ImportBatch(
            project_id=project.id,
            file_name="delay.xlsx",
            sheet_name="消防单位",
            status="published",
            data_date=date(2026, 5, 12),
            imported_count=3,
            baseline_plan_id=baseline.id,
        )
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="轻微项",
                    building="A1",
                    floor="2层",
                    discipline="消防",
                    system_name="喷淋",
                    actual_percent=80,
                        planned_percent=83,
                        progress_deviation=-3,
                        status="slightly_delayed",
                            planned_start_date=date(2026, 5, 1),
                            planned_finish_date=date(2026, 5, 14),
                    ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="严重项",
                    building="A1",
                    floor="1层",
                    discipline="消防",
                    system_name="报警",
                    actual_percent=30,
                    planned_percent=60,
                        progress_deviation=-30,
                        status="seriously_delayed",
                        planned_start_date=date(2026, 5, 9),
                        planned_finish_date=date(2026, 5, 14),
                        extra_fields=json.dumps({"responsible_person": "李工"}, ensure_ascii=False),
                    ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="正常项",
                    building="A1",
                    floor="3层",
                    discipline="消防",
                    actual_percent=90,
                        planned_percent=80,
                        progress_deviation=10,
                        status="ahead",
                        planned_start_date=date(2026, 5, 1),
                        planned_finish_date=date(2026, 5, 21),
                    ),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(
            f"/api/projects/{project_id}/reports/delay-rectification-export",
            params={"batch_id": batch_id},
        )

    assert response.status_code == 200
    assert response.content[:2] == b"PK"
    workbook = load_workbook(BytesIO(response.content), read_only=True)
    try:
        sheet = workbook["整改清单"]
        rows = list(sheet.iter_rows(values_only=True))
        assert rows[0][0] == "滞后项整改清单"
        assert rows[1][0] == "项目名称"
        assert rows[3][0] == "批次绑定计划基线"
        assert rows[3][1] == "消防专项计划"
        header_index = next(index for index, row in enumerate(rows) if row[0] == "序号")
        headers = rows[header_index]
        assert "专业" in headers
        assert "楼栋" in headers
        assert "楼层" in headers
        assert "系统" in headers
        assert "施工项" in headers
        assert "滞后说明" in headers
        assert "整改建议" in headers
        assert "责任人" in headers
        assert "计划完成时间" in headers
        assert "复查结果" in headers
        data_rows = [row for row in rows[header_index + 1:] if isinstance(row[0], int)]
        assert data_rows[0][5] == "严重项"
        assert data_rows[0][8] == "-30.0%"
        assert data_rows[0][9] == "严重滞后"
        assert "【消防】A1 1层 严重项" in data_rows[0][10]
        assert "建议优先协调资源投入" in data_rows[0][11]
        assert data_rows[0][12] == "李工"
        assert data_rows[1][5] == "轻微项"
    finally:
        workbook.close()


def test_delay_rectification_export_handles_no_delayed_items() -> None:
    db = SessionLocal()
    try:
        project = Project(name="无滞后整改清单")
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="ok.xlsx", status="published", data_date=date(2026, 5, 12))
        db.add(batch)
        db.flush()
        db.add(ProgressItem(project_id=project.id, batch_id=batch.id, task_name="正常项", actual_percent=90, planned_percent=80, progress_deviation=10, status="ahead"))
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(f"/api/projects/{project_id}/reports/delay-rectification-export", params={"batch_id": batch_id})

    assert response.status_code == 200
    workbook = load_workbook(BytesIO(response.content), read_only=True)
    try:
        rows = list(workbook["整改清单"].iter_rows(values_only=True))
        assert any(row[0] == "当前筛选条件下暂无滞后项。" for row in rows)
    finally:
        workbook.close()


def test_baseline_plan_bindings_and_overview_use_bound_baseline() -> None:
    with TestClient(app) as client:
        project_id = client.post("/api/projects", json={"name": "基线管理测试"}).json()["id"]
        baseline_a = client.post(
            f"/api/projects/{project_id}/baseline-plans",
            json={"name": "基线A", "plan_type": "current", "description": "A", "baseline_date": "2026-05-01", "is_default": True, "is_active": True},
        ).json()
        baseline_b = client.post(
            f"/api/projects/{project_id}/baseline-plans",
            json={"name": "基线B", "plan_type": "adjusted", "description": "B", "baseline_date": "2026-05-08", "is_default": False, "is_active": True},
        ).json()
        assert baseline_a["is_default"] is True
        assert baseline_b["is_default"] is False

        batch = ImportBatch(
            project_id=project_id,
            file_name="base.xlsx",
            status="published",
            data_date=date(2026, 5, 12),
            imported_count=1,
            baseline_plan_id=baseline_b["id"],
        )
        db = SessionLocal()
        try:
            db.add(batch)
            db.flush()
            db.add(
                ProgressItem(
                    project_id=project_id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline_b["id"],
                    task_name="基线任务",
                    building="A1",
                    floor="1层",
                    discipline="消防",
                    actual_percent=60,
                    planned_percent=80,
                    progress_deviation=-20,
                    status="delayed",
                )
            )
            db.commit()
            batch_id = batch.id
        finally:
            db.close()

        overview = client.get(
            f"/api/projects/{project_id}/analytics/overview",
            params={"batch_id": batch_id},
        ).json()
        bound_batches = client.get(f"/api/projects/{project_id}/baseline-plans/{baseline_b['id']}/batches").json()
        baselines = client.get(f"/api/projects/{project_id}/baseline-plans").json()
        disabled = client.put(f"/api/projects/{project_id}/baseline-plans/{baseline_b['id']}", json={"is_active": False})
        report = client.get(
            f"/api/projects/{project_id}/reports/dashboard-export",
            params={"batch_id": batch_id, "baseline_plan_id": baseline_b["id"]},
        )

    assert overview["baseline_plan_id"] == baseline_b["id"]
    assert overview["baseline_plan_name"] == "基线B"
    assert next(item for item in baselines if item["id"] == baseline_b["id"])["bound_batch_count"] == 1
    assert next(item for item in baselines if item["id"] == baseline_b["id"])["latest_bound_batch_date"] == "2026-05-12"
    assert bound_batches[0]["file_name"] == "base.xlsx"
    assert bound_batches[0]["baseline_plan_name"] == "基线B"
    assert disabled.status_code == 200
    assert report.status_code == 200

    db = SessionLocal()
    try:
        historical_batch = db.get(ImportBatch, batch_id)
        assert historical_batch.baseline_plan_id == baseline_b["id"]
    finally:
        db.close()


def test_delay_rectification_export_handles_nullable_fields_and_legacy_delay_level() -> None:
    db = SessionLocal()
    try:
        project = Project(name='特殊/字符:项目?')
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="nullable.xlsx", status="published", data_date=date(2026, 5, 12))
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    task_name=None,
                    building=None,
                    floor=None,
                    discipline=None,
                    system_name=None,
                    actual_percent=0,
                    planned_percent=None,
                    progress_deviation=-12,
                    status=None,
                    planned_start_date=date(2026, 5, 1),
                    planned_finish_date=date(2026, 5, 13),
                    extra_fields="not-json",
                ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    task_name="旧枚举项",
                    building="A1",
                    floor="B1层",
                    discipline="消防",
                    actual_percent=0,
                    planned_percent=20,
                    progress_deviation=-12,
                    status="seriously_delay",
                    planned_start_date=date(2026, 5, 1),
                    planned_finish_date=date(2026, 5, 13),
                    extra_fields=None,
                ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    task_name="无偏差项",
                    progress_deviation=None,
                    actual_percent=None,
                    planned_percent=None,
                    status=None,
                ),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(
            f"/api/projects/{project_id}/reports/delay-rectification-export",
            params={"batch_id": batch_id, "delay_level": "严重滞后"},
        )

    assert response.status_code == 200
    assert response.content[:2] == b"PK"
    assert "特殊_字符_项目_滞后项整改清单_严重滞后_2026-05-12.xlsx" in unquote(response.headers["content-disposition"])
    workbook = load_workbook(BytesIO(response.content), read_only=True)
    try:
        rows = list(workbook["整改清单"].iter_rows(values_only=True))
        data_rows = [row for row in rows if isinstance(row[0], int)]
        assert len(data_rows) == 2
        assert data_rows[0][1] == "未填写专业"
        assert data_rows[0][2] == "未填写楼栋"
        assert data_rows[0][3] == "未填写楼层"
        assert data_rows[0][6] == "0.0%"
        assert data_rows[0][7] == "91.7%"
        assert data_rows[0][9] == "严重滞后"
        assert data_rows[0][12] in ("", None)
        assert data_rows[1][9] == "严重滞后"
        assert all(row[5] != "无偏差项" for row in data_rows)
    finally:
        workbook.close()


def test_delay_rectification_export_filters_and_records_history() -> None:
    db = SessionLocal()
    try:
        project = Project(name="筛选整改清单项目")
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="delay.xlsx", status="published", data_date=date(2026, 5, 12), imported_count=4)
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="消防A1", building="A1", floor="1层", discipline="消防", actual_percent=30, planned_percent=60, progress_deviation=-30, status="seriously_delayed", planned_start_date=date(2026, 5, 9), planned_finish_date=date(2026, 5, 14)),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="消防A2", building="A2", floor="2层", discipline="消防", actual_percent=60, planned_percent=66, progress_deviation=-6, status="delayed", planned_start_date=date(2026, 5, 1), planned_finish_date=date(2026, 5, 18)),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="电气A1", building="A1", floor="1层", discipline="电气", actual_percent=81, planned_percent=84, progress_deviation=-3, status="slightly_delayed", planned_start_date=date(2026, 5, 1), planned_finish_date=date(2026, 5, 14)),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="暖通A1", building="A1", floor="3层", discipline="暖通", actual_percent=90, planned_percent=80, progress_deviation=10, status="ahead", planned_start_date=date(2026, 5, 1), planned_finish_date=date(2026, 5, 21)),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    filter_cases = [
        ({"discipline": "消防"}, 2, 1, "消防"),
        ({"building": "A1"}, 2, 2, "A1"),
        ({"floor": "1层"}, 2, 3, "1层"),
        ({"delay_level": "seriously_delayed"}, 1, 9, "严重滞后"),
    ]
    with TestClient(app) as client:
        for params, expected_count, column_index, expected_value in filter_cases:
            response = client.get(
                f"/api/projects/{project_id}/reports/delay-rectification-export",
                params={"batch_id": batch_id, **params},
            )
            assert response.status_code == 200
            workbook = load_workbook(BytesIO(response.content), read_only=True)
            try:
                rows = list(workbook["整改清单"].iter_rows(values_only=True))
                data_rows = [row for row in rows if isinstance(row[0], int)]
                assert len(data_rows) == expected_count
                assert all(row[column_index] == expected_value for row in data_rows)
            finally:
                workbook.close()

        empty_response = client.get(
            f"/api/projects/{project_id}/reports/delay-rectification-export",
            params={"batch_id": batch_id, "discipline": "消防", "building": "A9"},
        )
        assert empty_response.status_code == 200
        empty_workbook = load_workbook(BytesIO(empty_response.content), read_only=True)
        try:
            rows = list(empty_workbook["整改清单"].iter_rows(values_only=True))
            assert any(row[0] == "当前筛选条件下暂无滞后项。" for row in rows)
        finally:
            empty_workbook.close()

    db = SessionLocal()
    try:
        records = db.query(ReportExportRecord).filter(ReportExportRecord.project_id == project_id).all()
        assert len(records) == 5
        assert all(record.report_type == "delay_rectification_excel" for record in records)
        assert all(record.file_name and record.file_path and record.batch_id == batch_id for record in records)
        assert all(record.data_date == date(2026, 5, 12) for record in records)
    finally:
        db.close()


def test_delay_rectification_report_type_alias_exports_and_records_history() -> None:
    db = SessionLocal()
    try:
        project = Project(name="统一类型整改清单")
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="delay.xlsx", status="published", data_date=date(2026, 5, 12))
        db.add(batch)
        db.flush()
        db.add(
            ProgressItem(
                project_id=project.id,
                batch_id=batch.id,
                task_name="喷淋支管",
                building="A1",
                floor="3层",
                discipline="消防",
                actual_percent=50,
                planned_percent=70,
                progress_deviation=-20,
                status="seriously_delayed",
                planned_start_date=date(2026, 5, 1),
                planned_finish_date=date(2026, 5, 15),
            )
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(
            f"/api/projects/{project_id}/reports/delay_rectification_excel",
            params={"batch_id": batch_id},
        )
        exports_response = client.get(f"/api/projects/{project_id}/reports/exports")

    assert response.status_code == 200
    assert response.content[:2] == b"PK"
    workbook = load_workbook(BytesIO(response.content), read_only=True)
    try:
        assert "整改清单" in workbook.sheetnames
        rows = list(workbook["整改清单"].iter_rows(values_only=True))
        assert rows[0][0] == "滞后项整改清单"
        assert any(row[5] == "喷淋支管" for row in rows if len(row) > 5)
    finally:
        workbook.close()

    assert exports_response.status_code == 200
    exports = exports_response.json()
    assert len(exports) == 1
    assert exports[0]["report_type"] == "delay_rectification_excel"
    assert exports[0]["data_date"] == "2026-05-12"


def test_legacy_delay_rectification_report_types_export_without_not_found() -> None:
    db = SessionLocal()
    try:
        project = Project(name="旧类型整改清单")
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="delay.xlsx", status="published", data_date=date(2026, 5, 12))
        db.add(batch)
        db.flush()
        db.add(
            ProgressItem(
                project_id=project.id,
                batch_id=batch.id,
                task_name="报警主机",
                discipline="消防",
                actual_percent=40,
                planned_percent=60,
                progress_deviation=-20,
                status="seriously_delayed",
            )
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        for legacy_type in ("delay_rectification", "rectification_excel", "delay_rectification_xlsx"):
            response = client.get(
                f"/api/projects/{project_id}/reports/{legacy_type}",
                params={"batch_id": batch_id},
            )
            assert response.status_code == 200
            assert response.content[:2] == b"PK"

    db = SessionLocal()
    try:
        records = db.query(ReportExportRecord).filter(ReportExportRecord.project_id == project_id).all()
        assert len(records) == 3
        assert all(record.report_type == "delay_rectification_excel" for record in records)
    finally:
        db.close()


def test_unknown_report_type_returns_structured_error() -> None:
    db = SessionLocal()
    try:
        project = Project(name="未知报表项目")
        db.add(project)
        db.commit()
        project_id = project.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(f"/api/projects/{project_id}/reports/not_registered_report")

    assert response.status_code == 404
    assert response.json()["detail"] == {
        "code": "REPORT_TYPE_NOT_FOUND",
        "message": "报表类型不存在或未注册。",
    }


def test_report_exports_empty_history_is_stable() -> None:
    db = SessionLocal()
    try:
        project = Project(name="空报表历史项目")
        db.add(project)
        db.commit()
        project_id = project.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(f"/api/projects/{project_id}/reports/exports")

    assert response.status_code == 200
    assert response.json() == []


def test_weekly_word_returns_project_not_found_code() -> None:
    with TestClient(app) as client:
        response = client.get("/api/projects/999999/reports/weekly-word")

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "PROJECT_NOT_FOUND"


def test_weekly_pdf_returns_pdf_and_records_history() -> None:
    db = SessionLocal()
    try:
        project = Project(name="PDF周报项目", project_type="测试")
        db.add(project)
        db.flush()
        baseline = BaselinePlan(project_id=project.id, name="当前计划", is_default=True)
        db.add(baseline)
        db.flush()
        batch = ImportBatch(
            project_id=project.id,
            file_name="pdf-weekly.xlsx",
            sheet_name="消防单位",
            status="published",
            data_date=date(2026, 5, 12),
            imported_count=2,
            warning_count=1,
            error_count=0,
            data_quality_score=95,
            baseline_plan_id=baseline.id,
        )
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="喷淋系统",
                    building="A1",
                    floor="3层",
                    discipline="消防",
                    system_name="喷淋",
                    unit="项",
                    actual_percent=58,
                    planned_percent=69,
                    progress_deviation=-11,
                    status="seriously_delayed",
                ),
                ProgressItem(
                    project_id=project.id,
                    batch_id=batch.id,
                    baseline_plan_id=baseline.id,
                    task_name="风管安装",
                    building="A2",
                    floor="1层",
                    discipline="暖通",
                    system_name="通风",
                    unit="米",
                    actual_percent=90,
                    planned_percent=80,
                    progress_deviation=10,
                    status="ahead",
                ),
            ]
        )
        db.add(ImportValidationIssue(batch_id=batch.id, level="warning", code="INVALID_DATE", message="bad date"))
        db.commit()
        project_id = project.id
        batch_id = batch.id
        baseline_id = baseline.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(
            f"/api/projects/{project_id}/reports/weekly-pdf",
            params={"batch_id": batch_id, "baseline_plan_id": baseline_id},
        )
        exports_response = client.get(f"/api/projects/{project_id}/reports/exports")

    assert response.status_code == 200
    assert response.content[:4] == b"%PDF"
    assert "application/pdf" in response.headers["content-type"]
    assert "PDF%E5%91%A8%E6%8A%A5%E9%A1%B9%E7%9B%AE_%E8%BF%9B%E5%BA%A6%E5%91%A8%E6%8A%A5_2026-05-12.pdf" in response.headers["content-disposition"]
    assert exports_response.status_code == 200
    exports = exports_response.json()
    assert exports[0]["report_type"] == "weekly_pdf"

    db = SessionLocal()
    try:
        record = db.query(ReportExportRecord).filter(ReportExportRecord.project_id == project_id).one()
        assert record.report_type == "weekly_pdf"
        assert record.file_name == "PDF周报项目_进度周报_2026-05-12.pdf"
        assert record.batch_id == batch_id
    finally:
        db.close()


def test_weekly_pdf_returns_clear_error_when_no_published_batch() -> None:
    db = SessionLocal()
    try:
        project = Project(name="无批次PDF周报项目")
        db.add(project)
        db.commit()
        project_id = project.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(f"/api/projects/{project_id}/reports/weekly-pdf")

    assert response.status_code == 404
    assert response.json()["detail"] == {"code": "NO_PUBLISHED_BATCH", "message": "当前暂无可导出数据。"}


def test_weekly_word_returns_clear_error_when_no_published_batch() -> None:
    db = SessionLocal()
    try:
        project = Project(name="无批次周报项目")
        db.add(project)
        db.commit()
        project_id = project.id
    finally:
        db.close()

    with TestClient(app) as client:
        response = client.get(f"/api/projects/{project_id}/reports/weekly-word")

    assert response.status_code == 404
    assert response.json()["detail"] == {"code": "NO_PUBLISHED_BATCH", "message": "当前暂无可导出数据。"}
