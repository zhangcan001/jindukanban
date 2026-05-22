from datetime import date
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook

from app.database import SessionLocal
from app.main import app
from app.models.import_batch import ImportBatch
from app.models.progress_item import ProgressItem
from app.models.project import Project
from app.services.analytics_service import aggregate_progress, available_calculation_methods


def _method_map(items):
    return {item["code"]: item for item in available_calculation_methods(items)}


def test_t031_recommends_weight_value_quantity_percent_and_task_methods() -> None:
    weighted = [ProgressItem(actual_percent=20, time_planned_percent=30, weight=0.4)]
    assert _method_map(weighted)["weighted_percent"]["recommended"] is True

    valued = [ProgressItem(actual_percent=20, value_amount=100)]
    assert _method_map(valued)["value_weighted_percent"]["recommended"] is True

    quantity = [ProgressItem(unit="米", total_quantity=100, cumulative_quantity=40)]
    assert _method_map(quantity)["quantity_percent"]["recommended"] is True

    percent = [ProgressItem(actual_percent=55)]
    assert _method_map(percent)["percent_average"]["recommended"] is True

    task_only = [ProgressItem(status="completed"), ProgressItem(status="normal")]
    assert _method_map(task_only)["task_average"]["recommended"] is True


def test_t031_mixed_units_keeps_quantity_selectable_but_not_recommended() -> None:
    items = [
        ProgressItem(unit="米", total_quantity=100, cumulative_quantity=40, actual_percent=40),
        ProgressItem(unit="台", total_quantity=2, cumulative_quantity=1, actual_percent=50),
    ]
    methods = _method_map(items)
    assert methods["quantity_percent"]["available"] is True
    assert methods["quantity_percent"]["recommended"] is False
    assert "直接汇总工程量可能失真" in methods["quantity_percent"]["warning"]
    assert methods["percent_average"]["recommended"] is True


def test_t031_progress_calculation_methods_are_correct() -> None:
    items = [
        ProgressItem(actual_percent=20, time_planned_percent=30, weight=0.25, total_quantity=100, cumulative_quantity=40, unit="米"),
        ProgressItem(actual_percent=60, time_planned_percent=50, weight=0.75, total_quantity=300, cumulative_quantity=180, unit="米"),
    ]
    weighted_actual, _, _ = aggregate_progress(items, None, "actual_percent", "weighted_percent")
    quantity_actual, _, _ = aggregate_progress(items, None, "actual_percent", "quantity_percent")
    percent_actual, _, _ = aggregate_progress(items, None, "actual_percent", "percent_average")
    assert weighted_actual == 50
    assert quantity_actual == 55
    assert percent_actual == 40


def test_t031_dashboard_and_reports_include_calculation_method_context() -> None:
    db = SessionLocal()
    try:
        project = Project(name="T031 动态统计口径")
        db.add(project)
        db.flush()
        batch = ImportBatch(project_id=project.id, file_name="t031.xlsx", status="published", data_date=date(2026, 5, 19))
        db.add(batch)
        db.flush()
        db.add_all(
            [
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="A", discipline="机电", unit="米", actual_percent=20, planned_percent=30, time_planned_percent=30, weight=0.4, extra_fields='{"权重": "40%"}'),
                ProgressItem(project_id=project.id, batch_id=batch.id, task_name="B", discipline="机电", unit="台", actual_percent=60, planned_percent=50, time_planned_percent=50, weight=0.6, extra_fields='{"权重": "60%"}'),
            ]
        )
        db.commit()
        project_id = project.id
        batch_id = batch.id
    finally:
        db.close()

    with TestClient(app) as client:
        overview = client.get(f"/api/projects/{project_id}/analytics/overview", params={"batch_id": batch_id})
        dashboard = client.get(f"/api/projects/{project_id}/reports/dashboard-export", params={"batch_id": batch_id})
        word = client.get(f"/api/projects/{project_id}/reports/weekly-word", params={"batch_id": batch_id})

    assert overview.status_code == 200
    payload = overview.json()
    assert payload["available_calculation_methods"][0]["code"] == "weighted_percent"
    assert any(item["recommended"] and item["code"] == "weighted_percent" for item in payload["available_calculation_methods"])

    workbook = load_workbook(BytesIO(dashboard.content), read_only=True)
    try:
        values = [cell for row in workbook["看板总览"].iter_rows(values_only=True) for cell in row]
        assert "统计口径" in values
        assert "权重统计" in values
        assert "统计口径说明" in values
    finally:
        workbook.close()

    document = Document(BytesIO(word.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "统计口径" in table_text
    assert "权重统计" in table_text
    assert "总体进度概况" in text
