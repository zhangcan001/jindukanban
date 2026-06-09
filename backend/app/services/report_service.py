from __future__ import annotations

from datetime import date, datetime, timedelta
from html import escape
import json
from pathlib import Path
import re
from uuid import uuid4

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models.import_batch import ImportBatch
from app.models.import_validation_issue import ImportValidationIssue
from app.models.progress_item import ProgressItem
from app.models.project import Project
from app.models.rectification_action_log import RectificationActionLog
from app.models.rectification_item import RectificationItem
from app.models.report_export_record import ReportExportRecord
from app.schemas.report import ReportConfig
from app.models.warning_record import WarningRecord
from app.models.warning_rule import WarningRule
from app.services.report_metadata import (
    DASHBOARD_EXCEL_TYPE,
    DELAY_RECTIFICATION_EXCEL_TYPE,
    REPORT_TYPES,
    STANDARD_EXCEL_REPORT_TYPES,
    WEEKLY_PDF_TYPE,
    WEEKLY_WORD_TYPE,
    resolve_report_config,
    serialize_report_config,
)
from app.services.analytics_service import (
    aggregate_progress,
    apply_time_based_progress,
    baseline_context,
    build_delay_message,
    calculation_method_description,
    delay_reference_date,
    delay_level_for_deviation,
    display_text,
    effective_baseline_plan,
    filter_items_by_baseline,
    get_published_batch,
    group_items,
    is_delay_eligible,
    item_units,
    list_items,
    quantity_sum,
    resolve_calculation_profile,
    sort_dimension_value,
    statistics_context,
)
from app.services.dashboard_plus_service import build_dashboard_plus
from app.services.progress_insight_service import generate_progress_insight
from app.services.ai_service import build_ai_insight_payload, fallback_insight_text, generate_ai_text_with_logging, project_ai_config
from app.services.warning_service import is_data_quality_warning_record

REPORT_DIR = Path(get_settings().export_dir)


def create_report(
    db: Session,
    project_id: int,
    report_type: str,
    batch_id: int | None,
    calculation_profile_id: int | None = None,
    baseline_plan_id: int | None = None,
    calculation_method: str | None = None,
) -> ReportExportRecord:
    if report_type not in STANDARD_EXCEL_REPORT_TYPES:
        raise ValueError(f"Unsupported report type: {report_type}")

    batch = get_published_batch(db, project_id, batch_id)
    if batch is None:
        raise LookupError("Published import batch not found")
    profile = resolve_calculation_profile(db, project_id, calculation_profile_id or batch.calculation_profile_id)
    items = apply_time_based_progress(list_items(db, project_id, batch.id), batch, profile)
    baseline = effective_baseline_plan(db, project_id, batch, baseline_plan_id)
    baseline_meta = baseline_context(db, project_id, batch, baseline_plan_id)
    items = _apply_plan_start_status(filter_items_by_baseline(items, baseline), batch)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = REPORT_TYPES[report_type]["label"][:31]
    _fill_report_context(sheet, batch, baseline_meta, items, profile, calculation_method)
    if report_type == "overview":
        _fill_overview(sheet, batch, items, profile, calculation_method)
    elif report_type == "delayed-ranking":
        _fill_delayed_ranking(sheet, items)
    elif report_type == "discipline-summary":
        _fill_discipline_summary(sheet, items, profile, calculation_method)
    elif report_type == "progress-items":
        _fill_progress_items(sheet, items)

    _style_workbook(workbook)
    project_dir = REPORT_DIR / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    file_name = f"{report_type}_{batch.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}.xlsx"
    file_path = project_dir / file_name
    workbook.save(file_path)

    record = ReportExportRecord(
        project_id=project_id,
        batch_id=batch.id,
        report_type=report_type,
        file_name=file_name,
        file_path=str(file_path),
        data_date=batch.data_date,
        exported_by="system",
    )
    db.add(record)
    db.flush()
    return record


def create_dashboard_export(
    db: Session,
    project: Project,
    batch_id: int | None,
    calculation_profile_id: int | None = None,
    baseline_plan_id: int | None = None,
    building: str | None = None,
    discipline: str | None = None,
    floor: str | None = None,
    delay_level: str | None = None,
    metric: str | None = None,
    calculation_method: str | None = None,
    construction_unit: str | None = None,
    system_name: str | None = None,
    scope: str | None = None,
    data_date: date | None = None,
    import_group_id: str | None = None,
    batch_ids: str | None = None,
) -> ReportExportRecord:
    batches = _resolve_report_batches(db, project.id, scope, batch_id, data_date, import_group_id, batch_ids)
    if not batches:
        raise LookupError("当前项目暂无已发布进度数据，无法导出看板。")
    batch = batches[-1]
    is_project_scope = (scope or "").strip().lower() == "project"

    profile = resolve_calculation_profile(db, project.id, calculation_profile_id or batch.calculation_profile_id)
    baseline = effective_baseline_plan(db, project.id, batch, baseline_plan_id)
    baseline_meta = baseline_context(db, project.id, batch, baseline_plan_id)
    filters = {
        "construction_unit": _clean_filter(construction_unit),
        "building": _clean_filter(building),
        "floor": _clean_filter(floor),
        "discipline": _clean_filter(discipline),
        "system_name": _clean_filter(system_name),
        "delay_level": _clean_filter(delay_level),
    }
    items = _scoped_report_items(db, project, batches, baseline, profile)
    items = _filter_dashboard_scope_items(items, filters)

    workbook = Workbook()
    overview_sheet = workbook.active
    overview_sheet.title = "看板总览"
    _fill_dashboard_overview(overview_sheet, db, project, batch, items, profile, baseline_meta, calculation_method, filters, batches if is_project_scope else None)
    _fill_dashboard_discipline(workbook.create_sheet("专业进度统计"), items, profile, calculation_method)
    _fill_dashboard_floor(workbook.create_sheet("楼层进度统计"), items, profile, calculation_method)
    _fill_dashboard_building_floor(workbook.create_sheet("楼栋楼层统计"), items, profile, building, calculation_method)
    _fill_dashboard_delayed(workbook.create_sheet("滞后项清单"), items)
    _fill_dashboard_quality(workbook.create_sheet("数据质量与校验问题汇总"), db, batch)
    rectifications = _filter_dashboard_rectifications(_rectification_items(db, project.id, None, [row.id for row in batches] if is_project_scope else None), filters)
    _fill_dashboard_rectification_summary(workbook.create_sheet("整改闭环摘要"), rectifications)
    _fill_dashboard_rectification_detail(workbook.create_sheet("整改项明细"), rectifications, db)
    insight = generate_progress_insight(db, project.id, batch.id, calculation_profile_id, baseline_plan_id, building)
    _fill_dashboard_insight(workbook.create_sheet("进度分析说明"), insight)
    if is_project_scope:
        _fill_dashboard_plus_discipline(workbook.create_sheet("专业进度对比"), None)
        _fill_dashboard_plus_floor_discipline(workbook.create_sheet("楼层专业矩阵"), None)
        _fill_dashboard_plus_building_discipline(workbook.create_sheet("楼栋专业矩阵"), None)
        _fill_dashboard_plus_delay_distribution(workbook.create_sheet("滞后分布统计"), None)
    else:
        dashboard_plus = build_dashboard_plus(
            db,
            project.id,
            batch.id,
            calculation_profile_id,
            calculation_method,
            baseline_plan_id,
            building,
            discipline,
            floor,
            construction_unit,
            system_name,
            delay_level,
            metric,
        )
        _fill_dashboard_plus_discipline(workbook.create_sheet("专业进度对比"), dashboard_plus)
        _fill_dashboard_plus_floor_discipline(workbook.create_sheet("楼层专业矩阵"), dashboard_plus)
        _fill_dashboard_plus_building_discipline(workbook.create_sheet("楼栋专业矩阵"), dashboard_plus)
        _fill_dashboard_plus_delay_distribution(workbook.create_sheet("滞后分布统计"), dashboard_plus)
    _style_workbook(workbook)

    project_dir = REPORT_DIR / str(project.id)
    project_dir.mkdir(parents=True, exist_ok=True)
    data_date = batch.data_date or date.today()
    file_name = f"{_safe_file_part(project.name)}_进度看板_{data_date.isoformat()}.xlsx"
    file_path = project_dir / f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}_{file_name}"
    workbook.save(file_path)

    record = ReportExportRecord(
        project_id=project.id,
        batch_id=None if is_project_scope else batch.id,
        report_type=DASHBOARD_EXCEL_TYPE,
        file_name=file_name,
        file_path=str(file_path),
        data_date=batch.data_date,
        exported_by="system",
    )
    db.add(record)
    db.flush()
    return record


def create_weekly_word_report(
    db: Session,
    project: Project,
    batch_id: int | None,
    calculation_profile_id: int | None = None,
    baseline_plan_id: int | None = None,
    building: str | None = None,
    use_ai_text: bool = False,
    calculation_method: str | None = None,
) -> ReportExportRecord:
    batch, profile, baseline_meta, items, report_config = _weekly_report_context(db, project, batch_id, calculation_profile_id, baseline_plan_id, "Word 周报")

    document = Document()
    _setup_word_document(document)
    _fill_weekly_cover(document, project, batch, profile, baseline_meta)
    _fill_weekly_overview(document, db, project, batch, items, profile, calculation_method)
    _fill_weekly_discipline(document, items, profile, calculation_method)
    _fill_weekly_floor(document, items, profile, calculation_method)
    _fill_weekly_building_floor(document, items, profile, building, calculation_method)
    _fill_weekly_delayed(document, items, report_config.weekly_delayed_item_limit)
    if report_config.show_rectification_summary:
        _fill_weekly_rectification_summary(document, _rectification_items(db, project.id, batch.id))
    if report_config.show_data_quality_section:
        _fill_weekly_quality(document, db, batch)
    insight = generate_progress_insight(db, project.id, batch.id, calculation_profile_id, baseline_plan_id, building)
    ai_weekly_text = _weekly_ai_text(db, project, batch.id, calculation_profile_id, baseline_plan_id, building) if use_ai_text or report_config.use_ai_weekly_text else None
    _fill_weekly_insight(document, insight, ai_weekly_text)
    if report_config.include_advanced_chart_analysis:
        dashboard_plus = build_dashboard_plus(db, project.id, batch.id, calculation_profile_id, calculation_method, baseline_plan_id, building)
        _fill_weekly_dashboard_plus(document, dashboard_plus, items, profile, report_config.weekly_matrix_summary_limit, calculation_method)

    project_dir = REPORT_DIR / str(project.id)
    project_dir.mkdir(parents=True, exist_ok=True)
    data_date = batch.data_date or date.today()
    file_name = f"{_safe_file_part(project.name)}_进度周报_{data_date.isoformat()}.docx"
    file_path = project_dir / f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}_{file_name}"
    document.save(file_path)

    record = ReportExportRecord(
        project_id=project.id,
        batch_id=batch.id,
        report_type=WEEKLY_WORD_TYPE,
        file_name=file_name,
        file_path=str(file_path),
        data_date=batch.data_date,
        exported_by="system",
    )
    db.add(record)
    db.flush()
    return record


def create_weekly_pdf_report(
    db: Session,
    project: Project,
    batch_id: int | None,
    calculation_profile_id: int | None = None,
    baseline_plan_id: int | None = None,
    building: str | None = None,
    use_ai: bool = False,
    calculation_method: str | None = None,
) -> ReportExportRecord:
    batch, profile, baseline_meta, items, report_config = _weekly_report_context(db, project, batch_id, calculation_profile_id, baseline_plan_id, "PDF 周报")
    rectifications = _rectification_items(db, project.id, batch.id) if report_config.show_rectification_summary else []
    insight = generate_progress_insight(db, project.id, batch.id, calculation_profile_id, baseline_plan_id, building)
    dashboard_plus = build_dashboard_plus(db, project.id, batch.id, calculation_profile_id, calculation_method, baseline_plan_id, building) if report_config.include_advanced_chart_analysis else None

    project_dir = REPORT_DIR / str(project.id)
    project_dir.mkdir(parents=True, exist_ok=True)
    data_date = batch.data_date or date.today()
    file_name = f"{_safe_file_part(project.name)}_进度周报_{data_date.isoformat()}.pdf"
    file_path = project_dir / f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}_{file_name}"
    _build_weekly_pdf(
        file_path,
        db,
        project,
        batch,
        profile,
        baseline_meta,
        items,
        report_config,
        insight,
        rectifications,
        dashboard_plus,
        building,
        calculation_method,
    )

    record = ReportExportRecord(
        project_id=project.id,
        batch_id=batch.id,
        report_type=WEEKLY_PDF_TYPE,
        file_name=file_name,
        file_path=str(file_path),
        data_date=batch.data_date,
        exported_by="system",
    )
    db.add(record)
    db.flush()
    return record


def _weekly_report_context(
    db: Session,
    project: Project,
    batch_id: int | None,
    calculation_profile_id: int | None,
    baseline_plan_id: int | None,
    report_label: str,
):
    batch = get_published_batch(db, project.id, batch_id)
    if batch is None:
        raise LookupError(f"当前项目暂无已发布进度数据，无法导出 {report_label}。")
    profile = resolve_calculation_profile(db, project.id, calculation_profile_id or batch.calculation_profile_id)
    baseline = effective_baseline_plan(db, project.id, batch, baseline_plan_id)
    baseline_meta = baseline_context(db, project.id, batch, baseline_plan_id)
    items = _apply_plan_start_status(filter_items_by_baseline(apply_time_based_progress(list_items(db, project.id, batch.id), batch, profile), baseline), batch)
    report_config = resolve_report_config(project.report_config)
    return batch, profile, baseline_meta, items, report_config


def create_delay_rectification_export(
    db: Session,
    project: Project,
    batch_id: int | None,
    calculation_profile_id: int | None = None,
    baseline_plan_id: int | None = None,
    building: str | None = None,
    discipline: str | None = None,
    floor: str | None = None,
    delay_level: str | None = None,
    calculation_method: str | None = None,
) -> ReportExportRecord:
    batch = get_published_batch(db, project.id, batch_id)
    if batch is None:
        raise LookupError("当前项目暂无已发布进度数据，无法导出整改清单。")

    profile = resolve_calculation_profile(db, project.id, calculation_profile_id or batch.calculation_profile_id)
    baseline = effective_baseline_plan(db, project.id, batch, baseline_plan_id)
    baseline_meta = baseline_context(db, project.id, batch, baseline_plan_id)
    items = _apply_plan_start_status(filter_items_by_baseline(apply_time_based_progress(list_items(db, project.id, batch.id), batch, profile), baseline), batch)
    filters = {
        "discipline": _clean_filter(discipline),
        "building": _clean_filter(building),
        "floor": _clean_filter(floor),
        "delay_level": _clean_filter(delay_level),
    }
    items = _filter_delay_rectification_items(items, filters)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "整改清单"
    profile = resolve_calculation_profile(db, project.id, calculation_profile_id or batch.calculation_profile_id)
    _fill_delay_rectification_sheet(sheet, project, batch, items, filters, baseline_meta, profile, calculation_method)
    _style_workbook(workbook)

    project_dir = REPORT_DIR / str(project.id)
    project_dir.mkdir(parents=True, exist_ok=True)
    data_date = batch.data_date or date.today()
    filter_parts = [_safe_file_part(value) for key, value in filters.items() if key != "delay_level" and value]
    if filters["delay_level"]:
        filter_parts.append(_safe_file_part(_status_label(filters["delay_level"])))
    filter_suffix = f"_{'_'.join(filter_parts)}" if filter_parts else ""
    file_name = f"{_safe_file_part(project.name)}_滞后项整改清单{filter_suffix}_{data_date.isoformat()}.xlsx"
    file_path = project_dir / f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}_{file_name}"
    workbook.save(file_path)

    record = ReportExportRecord(
        project_id=project.id,
        batch_id=batch.id,
        report_type=DELAY_RECTIFICATION_EXCEL_TYPE,
        file_name=file_name,
        file_path=str(file_path),
        data_date=batch.data_date,
        exported_by="system",
    )
    db.add(record)
    db.flush()
    return record


def _baseline_display(value: object) -> str:
    return str(value) if value else "未配置计划基线"


def _baseline_consistency_text(meta: dict[str, object]) -> str:
    return "是" if meta.get("baseline_consistent") else "否"


def _fill_report_context(
    sheet,
    batch: ImportBatch,
    baseline_meta: dict[str, object],
    items: list[ProgressItem] | None = None,
    profile=None,
    calculation_method: str | None = None,
) -> None:
    sheet.append(["报表上下文", "值"])
    sheet.append(["批次 ID", batch.id])
    sheet.append(["数据日期", batch.data_date])
    sheet.append(["批次绑定计划基线", _baseline_display(baseline_meta.get("batch_bound_baseline_plan_name"))])
    sheet.append(["当前查看计划基线", _baseline_display(baseline_meta.get("current_view_baseline_plan_name"))])
    sheet.append(["是否与批次绑定基线一致", _baseline_consistency_text(baseline_meta)])
    if items is not None:
        stats = statistics_context(items, profile, calculation_method)
        sheet.append(["统计口径", stats.label])
        sheet.append(["推荐原因", stats.reason or "用户手动选择统计口径"])
        sheet.append(["是否混合单位", "是" if len(item_units(items)) > 1 else "否"])
        sheet.append(["单位列表", "、".join(item_units(items)) or "未识别"])
        sheet.append(["权重来源", stats.weight_source or "未使用权重字段"])
        sheet.append(["当前范围权重合计", _weight_percent_text(stats.weight_total)])
        sheet.append(["是否归一化", "是" if stats.is_normalized else "否"])
        sheet.append(["统计口径说明", stats.method_description])
    if baseline_meta.get("baseline_notice"):
        sheet.append(["基线提示", baseline_meta["baseline_notice"]])
    sheet.append([])


def _fill_dashboard_overview(
    sheet,
    db: Session,
    project: Project,
    batch: ImportBatch,
    items: list[ProgressItem],
    profile,
    baseline_meta: dict[str, object],
    calculation_method: str | None = None,
    filters: dict[str, str | None] | None = None,
    scoped_batches: list[ImportBatch] | None = None,
) -> None:
    actual_percent, _, _ = aggregate_progress(items, profile, "actual_percent", calculation_method)
    planned_percent, _, _ = aggregate_progress(items, profile, "planned_percent", calculation_method)
    stats = statistics_context(items, profile, calculation_method)
    deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
    warning_record_count = _task_warning_record_count(db, project.id, batch.id)
    sheet.append(["看板总览"])
    sheet.append(["指标", "值"])
    rows = [
        ("项目名称", project.name),
        ("项目状态", "已归档" if project.is_archived else "正常"),
        ("数据日期", batch.data_date or ""),
        ("当前导出范围", "项目级聚合" if scoped_batches else "单批次"),
        ("包含 Sheet", "、".join(row.sheet_name or f"批次 {row.id}" for row in scoped_batches) if scoped_batches else (batch.sheet_name or f"批次 {batch.id}")),
        ("包含批次", "、".join(str(row.id) for row in scoped_batches) if scoped_batches else str(batch.id)),
        ("导入批次", f"#{batch.id} / {batch.file_name}"),
        ("批次状态", "已冻结" if batch.is_frozen else "正常"),
        ("统计口径", stats.label),
        ("推荐原因", stats.reason or "用户手动选择统计口径"),
        ("是否混合单位", "是" if len(item_units(items)) > 1 else "否"),
        ("单位列表", "、".join(item_units(items)) or "未识别"),
        ("统计口径说明", stats.method_description),
        ("统计口径配置", profile.name if profile else "项目默认口径"),
        ("当前筛选条件", _dashboard_filter_text(filters or {})),
        ("权重来源", stats.weight_source or "未使用权重字段"),
        ("当前范围权重合计", _weight_percent_text(stats.weight_total)),
        ("当前范围任务数", len(items)),
        ("是否归一化", "是" if stats.is_normalized else "否"),
        ("当前范围对项目总进度贡献", _percent_text(stats.project_contribution_actual)),
        ("当前范围对应完成贡献", _percent_text(stats.project_contribution_planned)),
        ("批次绑定计划基线", _baseline_display(baseline_meta.get("batch_bound_baseline_plan_name"))),
        ("当前查看计划基线", _baseline_display(baseline_meta.get("current_view_baseline_plan_name"))),
        ("是否与批次绑定基线一致", _baseline_consistency_text(baseline_meta)),
        ("基线提示", baseline_meta.get("baseline_notice") or ""),
        ("实际进度", _percent_text(actual_percent)),
        ("应完成进度", _percent_text(planned_percent)),
        ("进度偏差", _signed_percent_text(deviation)),
        ("任务数量", len({item.task_id for item in items if item.task_id is not None}) or len(items)),
        ("数据质量评分", batch.data_quality_score if batch.data_quality_score is not None else "-"),
        ("warning 数量", batch.warning_count or 0),
        ("error 数量", batch.error_count or 0),
        ("当前预警记录数", warning_record_count or 0),
        ("导出时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    for row in rows:
        sheet.append([_display_value(row[0]), _display_value(row[1], "")])


def _fill_dashboard_discipline(sheet, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    sheet.append(["专业进度统计"])
    sheet.append(["专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数", "单位情况", "备注"])
    for discipline, group in sorted(group_items(items, "discipline").items(), key=lambda row: sort_dimension_value("discipline", row[0])):
        actual_percent, unit_mixed, actual_warning = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, planned_unit_mixed, planned_warning = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        mixed = unit_mixed or planned_unit_mixed
        sheet.append([
            discipline,
            len(group),
            _percent_text(actual_percent),
            _percent_text(planned_percent),
            _signed_percent_text(deviation),
            _delayed_count(group),
            _unit_text(group, mixed),
            _unit_mixed_note() if mixed else (actual_warning or planned_warning or ""),
        ])


def _fill_dashboard_floor(sheet, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    sheet.append(["楼层进度统计"])
    sheet.append(["楼层", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数", "单位情况"])
    for floor, group in sorted(group_items(items, "floor").items(), key=lambda row: sort_dimension_value("floor", row[0])):
        actual_percent, unit_mixed, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, planned_unit_mixed, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        sheet.append([
            floor,
            len(group),
            _percent_text(actual_percent),
            _percent_text(planned_percent),
            _signed_percent_text(deviation),
            _delayed_count(group),
            _unit_text(group, unit_mixed or planned_unit_mixed),
        ])


def _fill_dashboard_building_floor(sheet, items: list[ProgressItem], profile, selected_building: str | None, calculation_method: str | None = None) -> None:
    sheet.append(["楼栋楼层统计"])
    sheet.append(["楼栋", "楼层", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数", "单位情况"])
    groups: dict[tuple[str, str], list[ProgressItem]] = {}
    for item in items:
        building = display_text(item.building, "未填写楼栋")
        if selected_building and building != selected_building:
            continue
        floor = display_text(item.floor, "未填写楼层")
        groups.setdefault((building, floor), []).append(item)
    for (building, floor), group in sorted(groups.items(), key=lambda row: (sort_dimension_value("building", row[0][0]), sort_dimension_value("floor", row[0][1]))):
        actual_percent, unit_mixed, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, planned_unit_mixed, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        sheet.append([
            building,
            floor,
            len(group),
            _percent_text(actual_percent),
            _percent_text(planned_percent),
            _signed_percent_text(deviation),
            _delayed_count(group),
            _unit_text(group, unit_mixed or planned_unit_mixed),
        ])


def _fill_dashboard_delayed(sheet, items: list[ProgressItem]) -> None:
    sheet.append(["滞后项清单"])
    sheet.append(["专业", "楼栋", "楼层", "系统", "施工项", "实际完成率", "应完成率", "进度偏差", "状态", "滞后说明"])
    for item in _delayed_items(items):
        _, label = delay_level_for_deviation(item.progress_deviation)
        sheet.append([
            display_text(item.discipline, "未填写专业"),
            display_text(item.building, "未填写楼栋"),
            display_text(item.floor, "未填写楼层"),
            display_text(item.system_name, "未填写系统"),
            display_text(item.task_name, "未填写施工项"),
            _percent_text(item.actual_percent),
            _percent_text(item.planned_percent),
            _signed_percent_text(item.progress_deviation),
            label,
            build_delay_message(item),
        ])


def _fill_dashboard_quality(sheet, db: Session, batch: ImportBatch) -> None:
    sheet.append(["数据质量与校验问题汇总"])
    sheet.append(["类型", "数量", "说明"])
    issue_rows = db.execute(
        select(ImportValidationIssue.level, ImportValidationIssue.code, func.count(ImportValidationIssue.id))
        .where(ImportValidationIssue.batch_id == batch.id)
        .group_by(ImportValidationIssue.level, ImportValidationIssue.code)
        .order_by(ImportValidationIssue.level, ImportValidationIssue.code)
    ).all()
    low_quality_message = _low_quality_message(batch)
    sheet.append(["数据质量评分", batch.data_quality_score if batch.data_quality_score is not None else "-", low_quality_message or "当前批次数据质量评分未低于建议阈值。"])
    if not issue_rows and not (batch.warning_count or batch.error_count):
        sheet.append(["info", 0, "当前批次暂无校验问题。"])
        return
    sheet.append(["warning", batch.warning_count or 0, "导入校验 warning 数量"])
    sheet.append(["error", batch.error_count or 0, "导入校验 error 数量"])
    for level, code, count in issue_rows:
        sheet.append([code or level or "UNKNOWN", count, f"{level or 'unknown'} 校验问题"])


def _fill_dashboard_rectification_summary(sheet, items: list[RectificationItem]) -> None:
    summary = _rectification_summary(items)
    sheet.append(["整改闭环摘要"])
    sheet.append(["指标", "值"])
    rows = [
        ("全部整改项", summary["total"]),
        ("未开始", summary["open"]),
        ("整改中", summary["in_progress"]),
        ("已完成", summary["completed"]),
        ("已关闭", summary["closed"]),
        ("已忽略", summary["ignored"]),
        ("逾期整改项", summary["overdue"]),
        ("严重滞后整改项", summary["serious"]),
        ("本周新增", summary["new_this_week"]),
        ("本周关闭", summary["closed_this_week"]),
    ]
    for row in rows:
        sheet.append(list(row))


def _fill_dashboard_rectification_detail(sheet, items: list[RectificationItem], db: Session) -> None:
    sheet.append(["整改项明细"])
    sheet.append([
        "状态",
        "滞后等级",
        "专业",
        "楼栋",
        "楼层",
        "系统",
        "施工项",
        "问题描述",
        "责任人",
        "责任单位",
        "计划完成时间",
        "是否逾期",
        "复查结果",
        "备注",
        "来源",
        "创建时间",
        "更新时间",
        "关闭时间",
    ])
    for item in items:
        sheet.append([
            _rectification_status_label(item.status),
            _rectification_delay_label(item.delay_level),
            display_text(item.discipline, "未填写专业"),
            display_text(item.building, "未填写楼栋"),
            display_text(item.floor, "未填写楼层"),
            display_text(item.system_name, "未填写系统"),
            display_text(item.task_name, "未填写施工项"),
            item.issue_description or "",
            item.responsible_person or "",
            item.responsible_unit or "",
            item.planned_finish_date.isoformat() if item.planned_finish_date else "",
            "是" if _rectification_overdue(item) else "否",
            item.review_result or "",
            item.remark or "",
            _rectification_source_label(item.source_type),
            item.created_at.strftime("%Y-%m-%d %H:%M") if item.created_at else "",
            item.updated_at.strftime("%Y-%m-%d %H:%M") if item.updated_at else "",
            item.closed_at.strftime("%Y-%m-%d %H:%M") if item.closed_at else "",
        ])


def _fill_dashboard_insight(sheet, insight) -> None:
    sheet.append(["进度分析说明"])
    sheet.append(["类型", "内容"])
    rows = [
        ("总体进度说明", insight.overview_summary),
        ("主要滞后专业", insight.discipline_summary),
        ("主要滞后楼层", insight.floor_summary),
        ("楼栋楼层说明", insight.building_floor_summary),
        ("主要滞后施工项", insight.delay_summary),
        ("数据质量说明", insight.quality_summary),
        ("本期关注事项", "\n".join(insight.focus_points)),
        ("建议措施", "\n".join(insight.recommended_actions)),
    ]
    for row in rows:
        sheet.append(list(row))


def _fill_dashboard_plus_discipline(sheet, dashboard_plus) -> None:
    sheet.append(["专业进度对比"])
    if dashboard_plus is None:
        sheet.append(["当前导出范围", "项目级聚合"])
        sheet.append(["说明", "项目级聚合范围已在专业进度统计、楼层进度统计、楼栋楼层统计和滞后项清单中导出。"])
        return
    _append_dashboard_plus_filters(sheet, dashboard_plus.filters)
    rows = dashboard_plus.discipline_progress
    sheet.append(["专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数", "严重滞后任务数", "是否严重滞后", "单位情况"])
    if not rows:
        sheet.append(["暂无数据"])
        return
    for row in rows:
        sheet.append([
            row.discipline,
            row.task_count,
            _percent_text(row.actual_percent),
            _percent_text(row.planned_percent),
            _signed_percent_text(row.progress_deviation),
            row.delayed_count,
            row.seriously_delayed_count,
            "是" if row.is_seriously_delayed else "否",
            _unit_text_from_values(row.units, row.unit_mixed),
        ])


def _fill_dashboard_plus_floor_discipline(sheet, dashboard_plus) -> None:
    sheet.append(["楼层专业矩阵"])
    if dashboard_plus is None:
        sheet.append(["当前导出范围", "项目级聚合"])
        sheet.append(["说明", "项目级聚合范围已在楼层进度统计和楼栋楼层统计中导出。"])
        return
    _append_dashboard_plus_filters(sheet, dashboard_plus.filters)
    rows = dashboard_plus.floor_discipline_matrix
    sheet.append(["楼层", "专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数"])
    if not rows:
        sheet.append(["暂无数据"])
        return
    for row in rows:
        sheet.append([
            row.floor,
            row.discipline,
            row.task_count,
            _percent_text(row.actual_percent),
            _percent_text(getattr(row, "planned_percent", None)),
            _signed_percent_text(row.progress_deviation),
            row.delayed_count,
        ])


def _fill_dashboard_plus_building_discipline(sheet, dashboard_plus) -> None:
    sheet.append(["楼栋专业矩阵"])
    if dashboard_plus is None:
        sheet.append(["当前导出范围", "项目级聚合"])
        sheet.append(["说明", "项目级聚合范围已在楼栋楼层统计中导出。"])
        return
    _append_dashboard_plus_filters(sheet, dashboard_plus.filters)
    rows = dashboard_plus.building_discipline_matrix
    sheet.append(["楼栋", "专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数"])
    if not rows:
        sheet.append(["暂无数据"])
        return
    for row in rows:
        sheet.append([
            row.building,
            row.discipline,
            row.task_count,
            _percent_text(row.actual_percent),
            _percent_text(getattr(row, "planned_percent", None)),
            _signed_percent_text(getattr(row, "progress_deviation", None)),
            row.delayed_count,
        ])


def _fill_dashboard_plus_delay_distribution(sheet, dashboard_plus) -> None:
    if dashboard_plus is None:
        sheet.append(["滞后分布统计"])
        sheet.append(["当前导出范围", "项目级聚合"])
        sheet.append(["说明", "项目级聚合范围已在滞后项清单中导出。"])
        return
    distribution = dashboard_plus.delay_distribution
    sheet.append(["滞后分布统计"])
    _append_dashboard_plus_filters(sheet, dashboard_plus.filters)
    sheet.append(["按状态统计"])
    sheet.append(["状态", "数量"])
    for row in distribution.status_counts:
        sheet.append([row.status_label, row.count])
    sheet.append([])
    sheet.append(["按专业分组"])
    sheet.append(["专业", "严重滞后", "明显滞后", "轻微滞后", "正常", "超前", "滞后合计"])
    if not distribution.discipline_delay_counts:
        sheet.append(["暂无数据"])
        return
    for row in distribution.discipline_delay_counts:
        sheet.append([
            row.discipline,
            row.seriously_delayed_count,
            row.delayed_count,
            row.slightly_delayed_count,
            row.normal_count,
            row.ahead_count,
            row.total_delayed_count,
        ])


def _append_dashboard_plus_filters(sheet, filters: dict[str, str | None]) -> None:
    labels = {
        "discipline": "专业",
        "building": "楼栋",
        "floor": "楼层",
        "delay_level": "滞后等级",
        "metric": "指标类型",
    }
    metric_labels = {
        "actual_percent": "实际进度",
        "planned_percent": "应完成进度",
        "progress_deviation": "进度偏差",
        "task_count": "任务数",
    }
    sheet.append(["当前筛选条件"])
    for key in ("discipline", "building", "floor", "delay_level", "metric"):
        value = filters.get(key)
        if key == "delay_level" and value:
            value = _status_label(value)
        if key == "metric" and value:
            value = metric_labels.get(value, value)
        sheet.append([labels[key], value or "全部"])
    sheet.append([])


def _fill_delay_rectification_sheet(
    sheet,
    project: Project,
    batch: ImportBatch,
    items: list[ProgressItem],
    filters: dict[str, str | None],
    baseline_meta: dict[str, object],
    profile=None,
    calculation_method: str | None = None,
) -> None:
    data_date = (batch.data_date or date.today()).isoformat()
    stats = statistics_context(items, profile, calculation_method)
    units = item_units(items)
    sheet.append(["滞后项整改清单"])
    sheet.append(["项目名称", _display_value(project.name)])
    sheet.append(["数据日期", data_date])
    sheet.append(["批次绑定计划基线", _baseline_display(baseline_meta.get("batch_bound_baseline_plan_name"))])
    sheet.append(["当前查看计划基线", _baseline_display(baseline_meta.get("current_view_baseline_plan_name"))])
    sheet.append(["是否与批次绑定基线一致", _baseline_consistency_text(baseline_meta)])
    sheet.append(["当前统计口径", stats.label])
    sheet.append(["推荐原因", stats.reason or "用户手动选择统计口径"])
    sheet.append(["是否混合单位", "是" if len(units) > 1 else "否"])
    sheet.append(["单位列表", "、".join(units) or "未识别"])
    sheet.append(["权重来源", stats.weight_source or "未使用权重字段"])
    sheet.append(["统计口径说明", stats.method_description])
    if baseline_meta.get("baseline_notice"):
        sheet.append(["基线提示", baseline_meta["baseline_notice"]])
    sheet.append(["筛选专业", _display_value(filters.get("discipline"), "全部")])
    sheet.append(["筛选楼栋", _display_value(filters.get("building"), "全部")])
    sheet.append(["筛选楼层", _display_value(filters.get("floor"), "全部")])
    sheet.append(["筛选滞后等级", _status_label(filters.get("delay_level")) if filters.get("delay_level") else "全部"])
    sheet.append(["导出时间", datetime.now().strftime("%Y-%m-%d %H:%M")])
    sheet.append([])
    sheet.append([
        "序号",
        "专业",
        "楼栋",
        "楼层",
        "系统",
        "施工项",
        "实际完成率",
        "应完成率",
        "偏差",
        "滞后等级",
        "滞后说明",
        "整改建议",
        "责任人",
        "计划完成时间",
        "复查结果",
        "备注",
    ])
    delayed = _delayed_items(items)
    if not delayed:
        sheet.append(["当前筛选条件下暂无滞后项。"])
        return
    for index, item in enumerate(delayed, start=1):
        level = _status_label(item.status)
        sheet.append([
            index,
            display_text(item.discipline, "未填写专业"),
            display_text(item.building, "未填写楼栋"),
            display_text(item.floor, "未填写楼层"),
            display_text(item.system_name, "未填写系统"),
            display_text(item.task_name, "未填写施工项"),
            _percent_text(item.actual_percent),
            _percent_text(item.planned_percent),
            _signed_percent_text(item.progress_deviation),
            level,
            build_delay_message(item),
            _rectification_suggestion(item.status),
            _responsible_person(item),
            "",
            "",
            "",
        ])


def _setup_word_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    export_time = datetime.now().strftime("%Y-%m-%d %H:%M")
    header = section.header.paragraphs[0]
    header.text = "工程进度周报"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.text = f"导出时间：{export_time}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fill_weekly_cover(document: Document, project: Project, batch: ImportBatch, profile, baseline_meta: dict[str, object]) -> None:
    title = document.add_heading("工程进度周报", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(22)
    subtitle = document.add_paragraph("项目进度管理系统自动生成")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    data_date = batch.data_date or date.today()
    rows = [
        ("项目名称", project.name),
        ("项目状态", "已归档" if project.is_archived else "正常"),
        ("数据日期", data_date.isoformat()),
        ("导出时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("统计口径配置", profile.name if profile else "项目默认口径"),
        ("批次状态", "已冻结" if batch.is_frozen else "正常"),
        ("批次绑定计划基线", _baseline_display(baseline_meta.get("batch_bound_baseline_plan_name"))),
        ("当前查看计划基线", _baseline_display(baseline_meta.get("current_view_baseline_plan_name"))),
        ("是否与批次绑定基线一致", _baseline_consistency_text(baseline_meta)),
        ("数据来源批次", f"#{batch.id} / {batch.file_name}"),
        ("任务数量", str(batch.imported_count or "-")),
        ("数据质量评分", str(batch.data_quality_score if batch.data_quality_score is not None else "-")),
    ]
    if baseline_meta.get("baseline_notice"):
        rows.append(("基线提示", str(baseline_meta["baseline_notice"])))
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = _display_value(value)
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    _style_word_table(table)
    document.add_paragraph()


def _fill_weekly_overview(document: Document, db: Session, project: Project, batch: ImportBatch, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    document.add_heading("一、总体进度概况", level=1)
    actual_percent, _, _ = aggregate_progress(items, profile, "actual_percent", calculation_method)
    planned_percent, _, _ = aggregate_progress(items, profile, "planned_percent", calculation_method)
    stats = statistics_context(items, profile, calculation_method)
    deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
    warning_record_count = _task_warning_record_count(db, project.id, batch.id)
    data_date = (batch.data_date or date.today()).isoformat()
    task_count = len({item.task_id for item in items if item.task_id is not None}) or len(items)
    if planned_percent is None or deviation is None:
        document.add_paragraph(
            f"截至 {data_date}，本项目实际进度为 {_percent_text(actual_percent)}。"
            "当前批次缺少应完成进度字段，暂无法计算计划偏差和滞后程度，建议后续补充应完成率或计划完成量字段。"
        )
    else:
        paragraph_text = (
            f"截至 {data_date}，本项目实际进度为 {_percent_text(actual_percent)}，应完成进度为 {_percent_text(planned_percent)}，"
            f"整体偏差为 {_signed_number_text(deviation)} 个百分点。当前批次共包含 {task_count} 项进度任务，"
            f"数据质量评分为 {batch.data_quality_score if batch.data_quality_score is not None else '-'} 分。"
        )
        if deviation <= -10:
            paragraph_text += "整体进度明显滞后，建议重点关注严重滞后专业、楼栋、楼层及施工项。"
        elif deviation <= -5:
            paragraph_text += "整体进度存在一定滞后，建议持续跟踪主要滞后项并落实整改措施。"
        elif deviation < 0:
            paragraph_text += "整体进度轻微滞后，建议关注局部滞后任务，防止滞后扩大。"
        else:
            paragraph_text += "整体进度总体可控，建议保持当前推进节奏。"
        document.add_paragraph(paragraph_text)
    document.add_heading("重点指标表", level=2)
    table = _add_word_table(
        document,
        ["指标", "值"],
        [
            ["统计口径", stats.label],
            ["推荐原因", stats.reason or "用户手动选择统计口径"],
            ["是否混合单位", "是" if len(item_units(items)) > 1 else "否"],
            ["单位列表", "、".join(item_units(items)) or "未识别"],
            ["统计口径说明", stats.method_description or "-"],
            ["权重来源", stats.weight_source or "未使用权重字段"],
            ["当前范围权重合计", _weight_percent_text(stats.weight_total)],
            ["是否归一化", "是" if stats.is_normalized else "否"],
            ["当前范围对项目总进度贡献", _percent_text(stats.project_contribution_actual)],
            ["实际进度", _percent_text(actual_percent)],
            ["应完成进度", _percent_text(planned_percent)],
            ["进度偏差", _signed_percent_text(deviation)],
            ["任务数量", str(task_count)],
            ["数据质量评分", str(batch.data_quality_score if batch.data_quality_score is not None else "-")],
            ["Warning", str(batch.warning_count or 0)],
            ["Error", str(batch.error_count or 0)],
            ["滞后项数量", str(_delayed_count(items))],
        ],
    )
    table.columns[0].width = Inches(1.6)


def _fill_weekly_discipline(document: Document, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    document.add_heading("二、分专业进度情况", level=1)
    rows = []
    delayed_candidates: list[tuple[float, str]] = []
    for discipline, group in sorted(group_items(items, "discipline").items(), key=lambda row: sort_dimension_value("discipline", row[0])):
        actual_percent, unit_mixed, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, planned_unit_mixed, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        mixed = unit_mixed or planned_unit_mixed
        if deviation is not None and deviation < 0:
            delayed_candidates.append((deviation, str(discipline)))
        rows.append([
            str(discipline),
            str(len(group)),
            _percent_text(actual_percent),
            _percent_text(planned_percent),
            _signed_percent_text(deviation),
            str(_delayed_count(group)),
            "当前分组包含多种单位，系统按当前统计口径计算，不直接汇总数量。" if mixed else "",
        ])
    _add_word_table(document, ["专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数", "备注"], rows)
    delayed_candidates.sort(key=lambda item: item[0])
    if delayed_candidates:
        names = "、".join(name for _, name in delayed_candidates[:3])
        document.add_paragraph(f"当前滞后较明显的专业主要包括：{names}。")
    else:
        document.add_paragraph("当前批次暂无可判断的专业滞后信息。")


def _fill_weekly_floor(document: Document, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    document.add_heading("三、楼层进度情况", level=1)
    rows = []
    delayed_candidates: list[tuple[float, str]] = []
    for floor, group in sorted(group_items(items, "floor").items(), key=lambda row: sort_dimension_value("floor", row[0])):
        actual_percent, _, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, _, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        if deviation is not None and deviation < 0:
            delayed_candidates.append((deviation, str(floor)))
        rows.append([str(floor), str(len(group)), _percent_text(actual_percent), _percent_text(planned_percent), _signed_percent_text(deviation), str(_delayed_count(group))])
    _add_word_table(document, ["楼层", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数"], rows)
    if rows:
        delayed_candidates.sort(key=lambda item: item[0])
        if delayed_candidates:
            document.add_paragraph(f"当前滞后较明显的楼层主要包括：{'、'.join(name for _, name in delayed_candidates[:3])}。")
        else:
            document.add_paragraph("当前批次暂无可判断的楼层滞后信息。")
    else:
        document.add_paragraph("当前批次暂无楼层统计数据。")


def _fill_weekly_building_floor(document: Document, items: list[ProgressItem], profile, selected_building: str | None, calculation_method: str | None = None) -> None:
    document.add_heading("四、楼栋楼层进度情况", level=1)
    groups: dict[tuple[str, str], list[ProgressItem]] = {}
    for item in items:
        building = display_text(item.building, "未填写楼栋")
        if selected_building and building != selected_building:
            continue
        floor = display_text(item.floor, "未填写楼层")
        groups.setdefault((building, floor), []).append(item)
    rows = []
    for (building, floor), group in sorted(groups.items(), key=lambda row: (sort_dimension_value("building", row[0][0]), sort_dimension_value("floor", row[0][1]))):
        actual_percent, _, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, _, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        rows.append([building, floor, str(len(group)), _percent_text(actual_percent), _percent_text(planned_percent), _signed_percent_text(deviation), str(_delayed_count(group))])
    _add_word_table(document, ["楼栋", "楼层", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数"], rows)


def _fill_weekly_delayed(document: Document, items: list[ProgressItem], limit: int = 30) -> None:
    document.add_heading("五、主要滞后项", level=1)
    delayed = _delayed_items(items)
    serious_count = sum(1 for item in delayed if _delay_level_value(item) == "seriously_delayed")
    obvious_count = sum(1 for item in delayed if _delay_level_value(item) == "delayed")
    slight_count = sum(1 for item in delayed if _delay_level_value(item) == "slightly_delayed")
    document.add_paragraph(f"本期主要滞后项共 {len(delayed)} 项，其中严重滞后 {serious_count} 项，明显滞后 {obvious_count} 项，轻微滞后 {slight_count} 项。")
    if not delayed:
        document.add_paragraph("当前批次暂无滞后项。")
        return
    rows = []
    for index, item in enumerate(delayed[: max(0, limit)], start=1):
        rows.append([
            str(index),
            display_text(item.discipline, "未填写专业"),
            display_text(item.building, "未填写楼栋"),
            display_text(item.floor, "未填写楼层"),
            display_text(item.system_name, "未填写系统"),
            display_text(item.task_name, "未填写施工项"),
            _percent_text(item.actual_percent),
            _percent_text(item.planned_percent),
            _signed_percent_text(item.progress_deviation),
            _status_label(item.status),
            build_delay_message(item),
        ])
    _add_word_table(document, ["序号", "专业", "楼栋", "楼层", "系统", "施工项", "实际%", "应完成%", "偏差", "状态", "滞后说明"], rows)


def _fill_weekly_rectification_summary(document: Document, items: list[RectificationItem]) -> None:
    document.add_heading("整改闭环摘要", level=1)
    summary = _rectification_summary(items)
    if summary["total"] == 0:
        document.add_paragraph("当前项目暂无整改闭环记录。")
    else:
        open_count = summary["open"] + summary["in_progress"] + summary["completed"]
        document.add_paragraph(
            f"当前项目共有 {summary['total']} 条整改项，其中未关闭 {open_count} 条，逾期 {summary['overdue']} 条。"
            "建议优先跟踪逾期及严重滞后整改项，确保责任人、完成时间和复查结果闭环。"
        )
    _add_word_table(
        document,
        ["指标", "数量"],
        [
            ["全部整改项", str(summary["total"])],
            ["未关闭整改项", str(summary["open"] + summary["in_progress"] + summary["completed"])],
            ["整改中", str(summary["in_progress"])],
            ["已完成", str(summary["completed"])],
            ["已关闭", str(summary["closed"])],
            ["逾期整改项", str(summary["overdue"])],
            ["严重滞后整改项", str(summary["serious"])],
        ],
    )


def _fill_weekly_quality(document: Document, db: Session, batch: ImportBatch) -> None:
    document.add_heading("六、数据质量与校验问题", level=1)
    document.add_paragraph(
        f"数据质量评分：{batch.data_quality_score if batch.data_quality_score is not None else '-'}；"
        f"warning 数量：{batch.warning_count or 0}；error 数量：{batch.error_count or 0}。"
    )
    low_quality_message = _low_quality_message(batch)
    if low_quality_message:
        document.add_paragraph(low_quality_message)
    issue_rows = db.execute(
        select(ImportValidationIssue.level, ImportValidationIssue.code, func.count(ImportValidationIssue.id))
        .where(ImportValidationIssue.batch_id == batch.id)
        .group_by(ImportValidationIssue.level, ImportValidationIssue.code)
        .order_by(ImportValidationIssue.level, ImportValidationIssue.code)
    ).all()
    if not issue_rows and not (batch.warning_count or batch.error_count):
        document.add_paragraph("当前批次暂无校验问题。")
        return
    rows = [[code or level or "UNKNOWN", str(count), _issue_description(code, level)] for level, code, count in issue_rows]
    if not rows:
        rows = [["warning", str(batch.warning_count or 0), "导入校验 warning 数量"], ["error", str(batch.error_count or 0), "导入校验 error 数量"]]
    _add_word_table(document, ["问题类型", "数量", "说明"], rows)


def _fill_weekly_insight(document: Document, insight, ai_text: str | None = None) -> None:
    document.add_heading("七、进度分析说明", level=1)
    if ai_text:
        document.add_paragraph("AI辅助生成：以下内容由 AI 辅助生成，请结合现场情况复核。")
        document.add_paragraph(ai_text)
        return
    sections = [
        ("总体进度说明", insight.overview_summary),
        ("主要滞后专业", insight.discipline_summary),
        ("主要滞后楼层", insight.floor_summary),
        ("楼栋楼层说明", insight.building_floor_summary),
        ("主要滞后施工项", insight.delay_summary),
        ("数据质量说明", insight.quality_summary),
    ]
    for title, text in sections:
        document.add_heading(title, level=2)
        document.add_paragraph(_display_value(text, "暂无相关说明。"))
    document.add_heading("本期关注事项", level=2)
    for note in insight.focus_points[:5]:
        document.add_paragraph(_display_value(note), style="List Bullet")
    document.add_heading("建议措施", level=2)
    for action in insight.recommended_actions[:5]:
        document.add_paragraph(_display_value(action), style="List Bullet")


def _weekly_ai_text(
    db: Session,
    project: Project,
    batch_id: int,
    calculation_profile_id: int | None,
    baseline_plan_id: int | None,
    building: str | None,
) -> str | None:
    config = project_ai_config(project)
    if not config.enabled or not config.api_base_url or not config.api_key or not config.model:
        return None
    try:
        payload = build_ai_insight_payload(db, project.id, batch_id, calculation_profile_id, baseline_plan_id, building)
        text, error_message, source = generate_ai_text_with_logging(
            db,
            project.id,
            batch_id,
            config,
            "weekly_report_text",
            payload,
            fallback_insight_text(db, project.id, batch_id, calculation_profile_id, baseline_plan_id, building),
        )
        db.commit()
        return text if source == "ai" and text else None
    except Exception:
        return None


def _fill_weekly_dashboard_plus(document: Document, dashboard_plus, items: list[ProgressItem], profile, matrix_limit: int = 10, calculation_method: str | None = None) -> None:
    document.add_heading("八、进阶图表分析", level=1)
    document.add_paragraph("本节基于 Dashboard 进阶图表口径，提炼适合会议汇报的专业、楼层、楼栋和滞后分布摘要。")

    document.add_heading("专业进度对比摘要", level=2)
    discipline_rows = [
        [
            row.discipline,
            str(row.task_count),
            _percent_text(row.actual_percent),
            _percent_text(row.planned_percent),
            _signed_percent_text(row.progress_deviation),
            str(row.delayed_count),
        ]
        for row in dashboard_plus.discipline_progress
    ]
    if discipline_rows:
        _add_word_table(document, ["专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后数量"], discipline_rows)
    else:
        document.add_paragraph("当前批次暂无专业进度对比数据。")

    matrix_rows = _advanced_matrix_rows(items, profile, calculation_method)
    document.add_heading("楼层专业矩阵摘要", level=2)
    floor_rows = _top_matrix_rows(matrix_rows["floor"], max(0, matrix_limit), ("floor", "discipline"))
    if floor_rows:
        _add_word_table(document, ["楼层", "专业", "任务数", "实际进度", "应完成进度", "偏差", "滞后数量"], floor_rows)
    else:
        document.add_paragraph("当前批次暂无楼层专业矩阵数据。")

    document.add_heading("楼栋专业矩阵摘要", level=2)
    building_rows = _top_matrix_rows(matrix_rows["building"], max(0, matrix_limit), ("building", "discipline"))
    if building_rows:
        _add_word_table(document, ["楼栋", "专业", "任务数", "实际进度", "应完成进度", "偏差", "滞后数量"], building_rows)
    else:
        document.add_paragraph("当前批次暂无楼栋专业矩阵数据。")

    document.add_heading("滞后分布统计", level=2)
    delay_rows = [[row.status_label, str(row.count)] for row in dashboard_plus.delay_distribution.status_counts]
    if delay_rows:
        _add_word_table(document, ["状态", "数量"], delay_rows)
    else:
        document.add_paragraph("当前批次暂无滞后分布数据。")


def _build_weekly_pdf(
    file_path: Path,
    db: Session,
    project: Project,
    batch: ImportBatch,
    profile,
    baseline_meta: dict[str, object],
    items: list[ProgressItem],
    report_config: ReportConfig,
    insight,
    rectifications: list[RectificationItem],
    dashboard_plus,
    building: str | None,
    calculation_method: str | None = None,
) -> None:
    font_name = _register_pdf_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "WeeklyPdfTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.black,
    )
    subtitle_style = ParagraphStyle(
        "WeeklyPdfSubtitle",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle("WeeklyPdfHeading", parent=styles["Heading2"], fontName=font_name, fontSize=12, leading=16, spaceBefore=8, spaceAfter=6)
    body_style = ParagraphStyle("WeeklyPdfBody", parent=styles["BodyText"], fontName=font_name, fontSize=9, leading=13)
    small_style = ParagraphStyle("WeeklyPdfSmall", parent=styles["BodyText"], fontName=font_name, fontSize=8, leading=11)

    story: list[object] = []
    data_date = batch.data_date or date.today()
    story.append(Paragraph("工程进度周报", title_style))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("项目进度管理系统自动生成", subtitle_style))
    story.append(Spacer(1, 6 * mm))
    story.append(_pdf_key_value_table([
        ("项目名称", project.name),
        ("项目状态", "已归档" if project.is_archived else "正常"),
        ("数据日期", data_date.isoformat()),
        ("导出时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("统计口径配置", profile.name if profile else "项目默认口径"),
        ("批次状态", "已冻结" if batch.is_frozen else "正常"),
        ("批次绑定计划基线", _baseline_display(baseline_meta.get("batch_bound_baseline_plan_name"))),
        ("当前查看计划基线", _baseline_display(baseline_meta.get("current_view_baseline_plan_name"))),
        ("是否与批次绑定基线一致", _baseline_consistency_text(baseline_meta)),
        ("数据来源批次", f"#{batch.id} / {batch.file_name}"),
        ("任务数量", str(batch.imported_count or "-")),
        ("数据质量评分", str(batch.data_quality_score if batch.data_quality_score is not None else "-")),
    ], font_name, body_style))
    if baseline_meta.get("baseline_notice"):
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(f"基线提示：{_display_value(baseline_meta['baseline_notice'])}", small_style))

    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph("一、总体进度概况", heading_style))
    actual_percent, _, _ = aggregate_progress(items, profile, "actual_percent", calculation_method)
    planned_percent, _, _ = aggregate_progress(items, profile, "planned_percent", calculation_method)
    stats = statistics_context(items, profile, calculation_method)
    deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
    task_count = len({item.task_id for item in items if item.task_id is not None}) or len(items)
    overview_text = (
        f"截至 {data_date.isoformat()}，本项目实际进度为 {_percent_text(actual_percent)}，"
        f"应完成进度为 {_percent_text(planned_percent)}，整体偏差为 {_signed_number_text(deviation)} 个百分点。"
    )
    story.append(Paragraph(overview_text if planned_percent is not None and deviation is not None else f"截至 {data_date.isoformat()}，本项目实际进度为 {_percent_text(actual_percent)}。", body_style))
    story.append(_pdf_key_value_table([
        ("统计口径", stats.label),
        ("推荐原因", stats.reason or "用户手动选择统计口径"),
        ("是否混合单位", "是" if len(item_units(items)) > 1 else "否"),
        ("单位列表", "、".join(item_units(items)) or "未识别"),
        ("统计口径说明", stats.method_description or "-"),
        ("权重来源", stats.weight_source or "未使用权重字段"),
        ("当前范围权重合计", _weight_percent_text(stats.weight_total)),
        ("是否归一化", "是" if stats.is_normalized else "否"),
        ("当前范围对项目总进度贡献", _percent_text(stats.project_contribution_actual)),
        ("实际进度", _percent_text(actual_percent)),
        ("应完成进度", _percent_text(planned_percent)),
        ("进度偏差", _signed_percent_text(deviation)),
        ("任务数量", str(task_count)),
        ("数据质量评分", str(batch.data_quality_score if batch.data_quality_score is not None else "-")),
        ("Warning", str(batch.warning_count or 0)),
        ("Error", str(batch.error_count or 0)),
        ("滞后项数量", str(_delayed_count(items))),
    ], font_name, body_style))

    story.append(Paragraph("二、分专业进度", heading_style))
    story.append(_pdf_table(
        ["专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数", "备注"],
        _weekly_discipline_rows(items, profile),
        font_name,
        body_style,
    ))

    story.append(Paragraph("三、楼层进度", heading_style))
    story.append(_pdf_table(
        ["楼层", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数"],
        _weekly_floor_rows(items, profile),
        font_name,
        body_style,
    ))

    story.append(Paragraph("四、楼栋楼层进度", heading_style))
    story.append(_pdf_table(
        ["楼栋", "楼层", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后任务数"],
        _weekly_building_floor_rows(items, profile, building),
        font_name,
        body_style,
    ))

    story.append(Paragraph("五、主要滞后项", heading_style))
    delayed = _delayed_items(items)
    story.append(Paragraph(
        f"本期主要滞后项共 {len(delayed)} 项，其中严重滞后 {sum(1 for item in delayed if _delay_level_value(item) == 'seriously_delayed')} 项，"
        f"明显滞后 {sum(1 for item in delayed if _delay_level_value(item) == 'delayed')} 项，轻微滞后 {sum(1 for item in delayed if _delay_level_value(item) == 'slightly_delayed')} 项。",
        body_style,
    ))
    story.append(_pdf_table(
        ["序号", "专业", "楼栋", "楼层", "系统", "施工项", "实际%", "应完成%", "偏差", "状态", "滞后说明"],
        _weekly_delayed_rows(items, report_config.weekly_delayed_item_limit),
        font_name,
        body_style,
    ))

    story.append(Paragraph("六、数据质量与校验问题", heading_style))
    story.append(_pdf_quality_rows(db, batch, font_name, body_style))

    story.append(Paragraph("七、进度分析说明", heading_style))
    story.extend(_pdf_insight_story(insight, body_style, heading_style))

    if report_config.show_rectification_summary:
        story.append(Paragraph("八、整改闭环摘要", heading_style))
        story.append(_pdf_rectification_summary(rectifications, font_name, body_style))

    if report_config.include_advanced_chart_analysis and dashboard_plus is not None:
        story.append(Paragraph("九、进阶图表分析摘要", heading_style))
        story.append(_pdf_dashboard_plus_summary(dashboard_plus, font_name, body_style, report_config.weekly_matrix_summary_limit))

    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    doc.build(story, onFirstPage=lambda canvas, _doc: _pdf_page_frame(canvas, _doc, font_name), onLaterPages=lambda canvas, _doc: _pdf_page_frame(canvas, _doc, font_name))


def _register_pdf_fonts() -> str:
    candidates = [
        ("MicrosoftYaHei", r"C:\Windows\Fonts\msyh.ttc"),
        ("SimHei", r"C:\Windows\Fonts\simhei.ttf"),
        ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
        ("NotoSansCJK", r"C:\Windows\Fonts\NotoSansCJKsc-Regular.otf"),
    ]
    for font_name, font_path in candidates:
        path = Path(font_path)
        if path.exists():
            try:
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, str(path)))
                return font_name
            except Exception:
                continue
    return "Helvetica"


def _pdf_page_frame(canvas, doc, font_name: str) -> None:
    canvas.saveState()
    canvas.setFont(font_name, 9)
    canvas.drawCentredString(A4[0] / 2, 10 * mm, f"第 {canvas.getPageNumber()} 页")
    canvas.restoreState()


def _pdf_key_value_table(rows: list[tuple[str, object]], font_name: str, body_style: ParagraphStyle) -> Table:
    data = [[Paragraph(_escape_pdf_text(label), body_style), Paragraph(_escape_pdf_text(_display_value(value)), body_style)] for label, value in rows]
    table = Table(data, colWidths=[42 * mm, 132 * mm], repeatRows=0)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("LEADING", (0, 0), (-1, -1), 12),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def _pdf_table(headers: list[str], rows: list[list[object]], font_name: str, body_style: ParagraphStyle) -> Table:
    data = [[Paragraph(_escape_pdf_text(header), body_style) for header in headers]]
    for row in rows or [["-", "-", "-", "-", "-", "-", "-"][: len(headers)]]:
        data.append([Paragraph(_escape_pdf_text(_display_value(value)), body_style) for value in row])
    table = Table(data, colWidths=[174 * mm / max(1, len(headers))] * len(headers), repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("LEADING", (0, 0), (-1, -1), 11),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    return table


def _escape_pdf_text(value: object) -> str:
    return escape(_display_value(value))


def _weekly_discipline_rows(items: list[ProgressItem], profile) -> list[list[object]]:
    rows: list[list[object]] = []
    for discipline, group in sorted(group_items(items, "discipline").items(), key=lambda row: sort_dimension_value("discipline", row[0])):
        actual_percent, unit_mixed, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile))
        planned_percent, planned_unit_mixed, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        mixed = unit_mixed or planned_unit_mixed
        rows.append([str(discipline), str(len(group)), _percent_text(actual_percent), _percent_text(planned_percent), _signed_percent_text(deviation), str(_delayed_count(group)), "当前分组包含多种单位，系统按当前统计口径计算，不直接汇总数量。" if mixed else ""])
    return rows


def _weekly_floor_rows(items: list[ProgressItem], profile) -> list[list[object]]:
    rows: list[list[object]] = []
    for floor, group in sorted(group_items(items, "floor").items(), key=lambda row: sort_dimension_value("floor", row[0])):
        actual_percent, _, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile))
        planned_percent, _, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        rows.append([str(floor), str(len(group)), _percent_text(actual_percent), _percent_text(planned_percent), _signed_percent_text(deviation), str(_delayed_count(group))])
    return rows


def _weekly_building_floor_rows(items: list[ProgressItem], profile, selected_building: str | None) -> list[list[object]]:
    groups: dict[tuple[str, str], list[ProgressItem]] = {}
    for item in items:
        building_name = display_text(item.building, "未填写楼栋")
        if selected_building and building_name != selected_building:
            continue
        floor_name = display_text(item.floor, "未填写楼层")
        groups.setdefault((building_name, floor_name), []).append(item)
    rows: list[list[object]] = []
    for (building_name, floor_name), group in sorted(groups.items(), key=lambda row: (sort_dimension_value("building", row[0][0]), sort_dimension_value("floor", row[0][1]))):
        actual_percent, _, _ = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile))
        planned_percent, _, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        rows.append([building_name, floor_name, str(len(group)), _percent_text(actual_percent), _percent_text(planned_percent), _signed_percent_text(deviation), str(_delayed_count(group))])
    return rows


def _weekly_delayed_rows(items: list[ProgressItem], limit: int) -> list[list[object]]:
    rows: list[list[object]] = []
    for index, item in enumerate(_delayed_items(items)[: max(0, limit)], start=1):
        rows.append([str(index), display_text(item.discipline, "未填写专业"), display_text(item.building, "未填写楼栋"), display_text(item.floor, "未填写楼层"), display_text(item.system_name, "未填写系统"), display_text(item.task_name, "未填写施工项"), _percent_text(item.actual_percent), _percent_text(item.planned_percent), _signed_percent_text(item.progress_deviation), _status_label(item.status), build_delay_message(item)])
    if not rows:
        rows.append(["-", "-", "-", "-", "-", "-", "-", "-", "-", "-", "当前批次暂无滞后项。"])
    return rows


def _pdf_quality_rows(db: Session, batch: ImportBatch, font_name: str, body_style: ParagraphStyle) -> Table:
    issue_rows = db.execute(
        select(ImportValidationIssue.level, ImportValidationIssue.code, func.count(ImportValidationIssue.id))
        .where(ImportValidationIssue.batch_id == batch.id)
        .group_by(ImportValidationIssue.level, ImportValidationIssue.code)
        .order_by(ImportValidationIssue.level, ImportValidationIssue.code)
    ).all()
    rows = [["数据质量评分", str(batch.data_quality_score if batch.data_quality_score is not None else "-"), _low_quality_message(batch) or "当前批次数据质量评分未低于建议阈值。"]]
    rows.extend([[code or level or "UNKNOWN", str(count), _issue_description(code, level)] for level, code, count in issue_rows])
    if not rows:
        rows = [["warning", str(batch.warning_count or 0), "导入校验 warning 数量"], ["error", str(batch.error_count or 0), "导入校验 error 数量"]]
    return _pdf_table(["问题类型", "数量", "说明"], rows, font_name, body_style)


def _pdf_insight_story(insight, body_style: ParagraphStyle, heading_style: ParagraphStyle) -> list[object]:
    story: list[object] = []
    for title, text in [
        ("总体进度说明", insight.overview_summary),
        ("主要滞后专业", insight.discipline_summary),
        ("主要滞后楼层", insight.floor_summary),
        ("楼栋楼层说明", insight.building_floor_summary),
        ("主要滞后施工项", insight.delay_summary),
        ("数据质量说明", insight.quality_summary),
    ]:
        story.append(Paragraph(title, ParagraphStyle("WeeklyPdfSubHeading", parent=heading_style, fontSize=10, leading=13, spaceBefore=4, spaceAfter=2)))
        story.append(Paragraph(_display_value(text, "暂无相关说明。"), body_style))
    story.append(Paragraph("本期关注事项", ParagraphStyle("WeeklyPdfSubHeading2", parent=heading_style, fontSize=10, leading=13, spaceBefore=4, spaceAfter=2)))
    for note in insight.focus_points[:5]:
        story.append(Paragraph(f"• {_display_value(note)}", body_style))
    story.append(Paragraph("建议措施", ParagraphStyle("WeeklyPdfSubHeading3", parent=heading_style, fontSize=10, leading=13, spaceBefore=4, spaceAfter=2)))
    for action in insight.recommended_actions[:5]:
        story.append(Paragraph(f"• {_display_value(action)}", body_style))
    return story


def _pdf_rectification_summary(items: list[RectificationItem], font_name: str, body_style: ParagraphStyle) -> Table:
    summary = _rectification_summary(items)
    if summary["total"] == 0:
        return _pdf_key_value_table([("说明", "当前项目暂无整改闭环记录。")], font_name, body_style)
    return _pdf_key_value_table([
        ("全部整改项", summary["total"]),
        ("未关闭整改项", summary["open"] + summary["in_progress"] + summary["completed"]),
        ("整改中", summary["in_progress"]),
        ("已完成", summary["completed"]),
        ("已关闭", summary["closed"]),
        ("逾期整改项", summary["overdue"]),
        ("严重滞后整改项", summary["serious"]),
    ], font_name, body_style)


def _pdf_dashboard_plus_summary(dashboard_plus, font_name: str, body_style: ParagraphStyle, matrix_limit: int) -> Table:
    rows = []
    for row in dashboard_plus.discipline_progress[: max(0, matrix_limit)]:
        rows.append([row.discipline, row.task_count, _percent_text(row.actual_percent), _percent_text(row.planned_percent), _signed_percent_text(row.progress_deviation), row.delayed_count])
    if not rows:
        rows = [["-", "-", "-", "-", "-", "-"]]
    return _pdf_table(["专业", "任务数", "实际进度", "应完成进度", "进度偏差", "滞后数量"], rows, font_name, body_style)


def _advanced_matrix_rows(items: list[ProgressItem], profile, calculation_method: str | None = None) -> dict[str, list[dict[str, object]]]:
    return {
        "floor": _grouped_matrix_rows(items, profile, ("floor", "discipline"), calculation_method),
        "building": _grouped_matrix_rows(items, profile, ("building", "discipline"), calculation_method),
    }


def _grouped_matrix_rows(items: list[ProgressItem], profile, fields: tuple[str, str], calculation_method: str | None = None) -> list[dict[str, object]]:
    groups: dict[tuple[str, str], list[ProgressItem]] = {}
    for item in items:
        first = display_text(getattr(item, fields[0]), f"未填写{_dimension_label(fields[0])}")
        second = display_text(getattr(item, fields[1]), f"未填写{_dimension_label(fields[1])}")
        groups.setdefault((first, second), []).append(item)

    rows: list[dict[str, object]] = []
    algorithm = _group_algorithm(profile, calculation_method)
    for (first, second), group in groups.items():
        actual_percent, _, _ = aggregate_progress(group, profile, "actual_percent", algorithm)
        planned_percent, _, _ = aggregate_progress(group, profile, "planned_percent", algorithm)
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        rows.append(
            {
                fields[0]: first,
                fields[1]: second,
                "task_count": len(group),
                "actual_percent": actual_percent,
                "planned_percent": planned_percent,
                "progress_deviation": deviation,
                "delayed_count": _delayed_count(group),
            }
        )
    return rows


def _top_matrix_rows(rows: list[dict[str, object]], limit: int, fields: tuple[str, str]) -> list[list[str]]:
    def sort_key(row: dict[str, object]) -> tuple[float, int, str, str]:
        deviation = row.get("progress_deviation")
        sort_deviation = float(deviation) if deviation is not None else 9999.0
        return (
            sort_deviation,
            -int(row.get("delayed_count") or 0),
            str(row.get(fields[0]) or ""),
            str(row.get(fields[1]) or ""),
        )

    return [
        [
            _display_value(row.get(fields[0])),
            _display_value(row.get(fields[1])),
            str(row.get("task_count") or 0),
            _percent_text(row.get("actual_percent")),
            _percent_text(row.get("planned_percent")),
            _signed_percent_text(row.get("progress_deviation")),
            str(row.get("delayed_count") or 0),
        ]
        for row in sorted(rows, key=sort_key)[:limit]
    ]


def _dimension_label(field: str) -> str:
    return {"floor": "楼层", "building": "楼栋", "discipline": "专业"}.get(field, field)


def _add_word_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = _display_value(value)
    _style_word_table(table)
    document.add_paragraph()
    return table


def _style_word_table(table) -> None:
    header_fill = "D9E2F3"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_border(cell)
            if row_index == 0:
                _shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9.5)
                    if row_index == 0:
                        run.bold = True


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def _set_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), "A6A6A6")


def _fill_overview(sheet, batch: ImportBatch, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    actual_percent, unit_mixed, warning = aggregate_progress(items, profile, "actual_percent", calculation_method)
    planned_percent, _, _ = aggregate_progress(items, profile, "planned_percent", calculation_method)
    stats = statistics_context(items, profile, calculation_method)
    total_quantity, _, total_warning = quantity_sum(items, "total_quantity", profile)
    actual_quantity, _, _ = quantity_sum(items, "actual_quantity", profile)
    remaining_quantity, _, _ = quantity_sum(items, "remaining_quantity", profile)
    deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
    rows = [
        ("批次 ID", batch.id),
        ("数据日期", batch.data_date),
        ("导入明细数", batch.imported_count),
        ("统计口径", stats.label),
        ("推荐原因", stats.reason or "用户手动选择统计口径"),
        ("是否混合单位", "是" if len(item_units(items)) > 1 else "否"),
        ("单位列表", "、".join(item_units(items)) or "未识别"),
        ("统计口径说明", stats.method_description),
        ("权重来源", stats.weight_source or "未使用权重字段"),
        ("当前范围权重合计", _weight_percent_text(stats.weight_total)),
        ("是否归一化", "是" if stats.is_normalized else "否"),
        ("实际进度", actual_percent),
        ("应完成进度", planned_percent),
        ("进度偏差", deviation),
        ("总工程量", total_quantity),
        ("实际完成量", actual_quantity),
        ("剩余工程量", remaining_quantity),
        ("单位混杂", "是" if unit_mixed else "否"),
        ("提示", warning or total_warning or ""),
        ("数据质量评分", batch.data_quality_score),
    ]
    sheet.append(["指标", "值"])
    for row in rows:
        sheet.append(list(row))


def _fill_delayed_ranking(sheet, items: list[ProgressItem]) -> None:
    sheet.append(["专业", "楼栋", "楼层", "系统", "施工项", "实际完成率", "应完成率", "进度偏差", "状态", "滞后说明"])
    for item in _delayed_items(items):
        _, delay_label = delay_level_for_deviation(item.progress_deviation)
        sheet.append([
            display_text(item.discipline, "未填写专业"),
            display_text(item.building, "未填写楼栋"),
            display_text(item.floor, "未填写楼层"),
            display_text(item.system_name, "未填写系统"),
            display_text(item.task_name, "未填写施工项"),
            item.actual_percent,
            item.planned_percent,
            item.progress_deviation,
            delay_label,
            build_delay_message(item),
        ])


def _fill_discipline_summary(sheet, items: list[ProgressItem], profile, calculation_method: str | None = None) -> None:
    sheet.append(["专业", "任务数", "实际%", "应完成%", "偏差", "总工程量", "实际完成量", "单位混杂", "提示"])
    for discipline, group in group_items(items, "discipline").items():
        actual_percent, unit_mixed, warning = aggregate_progress(group, profile, "actual_percent", _group_algorithm(profile, calculation_method))
        planned_percent, _, _ = aggregate_progress(group, profile, "planned_percent", _group_algorithm(profile, calculation_method))
        deviation = round(actual_percent - planned_percent, 4) if actual_percent is not None and planned_percent is not None else None
        total_quantity, _, quantity_warning = quantity_sum(group, "total_quantity", profile)
        actual_quantity, _, _ = quantity_sum(group, "actual_quantity", profile)
        sheet.append([
            discipline,
            len(group),
            actual_percent,
            planned_percent,
            deviation,
            total_quantity,
            actual_quantity,
            "是" if unit_mixed else "否",
            warning or quantity_warning or "",
        ])


def _fill_progress_items(sheet, items: list[ProgressItem]) -> None:
    sheet.append([
        "任务",
        "WBS",
        "编码",
        "区域",
        "楼栋",
        "楼层",
        "专业",
        "系统",
        "单位",
        "总工程量",
        "计划完成量",
        "实际完成量",
        "实际%",
        "应完成%",
        "偏差",
        "状态",
        "人工修正",
        "备注",
    ])
    for item in items:
        sheet.append([
            item.task_name,
            item.wbs_code,
            item.task_code,
            item.area,
            item.building,
            item.floor,
            item.discipline,
            item.system_name,
            item.unit,
            item.total_quantity,
            item.planned_quantity,
            item.actual_quantity,
            item.actual_percent,
            item.planned_percent,
            item.progress_deviation,
            _progress_item_status_value(item),
            "是" if item.is_manually_edited else "否",
            item.remark,
        ])


def _group_algorithm(profile, calculation_method: str | None = None) -> str:
    return calculation_method or ((profile.group_algorithm if profile else "auto") or "auto")


def _apply_plan_start_status(items: list[ProgressItem], batch: ImportBatch) -> list[ProgressItem]:
    reference_date = delay_reference_date(batch)
    for item in items:
        if not is_delay_eligible(item, reference_date):
            setattr(item, "_not_started_by_plan_for_delay", True)
    return items


def _is_not_started_for_delay(item: ProgressItem) -> bool:
    return bool(getattr(item, "_not_started_by_plan_for_delay", False)) or item.status in {
        "not_started_by_plan",
        "missing_plan_dates",
        "invalid_plan_dates",
    }


def _progress_item_status_value(item: ProgressItem) -> str | None:
    return "not_started_by_plan" if _is_not_started_for_delay(item) else item.status


def _delayed_count(items: list[ProgressItem]) -> int:
    return sum(1 for item in items if not _is_not_started_for_delay(item) and item.progress_deviation is not None and item.progress_deviation < 0)


def _delayed_items(items: list[ProgressItem]) -> list[ProgressItem]:
    delayed = [item for item in items if not _is_not_started_for_delay(item) and item.progress_deviation is not None and item.progress_deviation < 0]
    delayed.sort(key=lambda item: item.progress_deviation or 0)
    return delayed


def _filter_delay_rectification_items(items: list[ProgressItem], filters: dict[str, str | None]) -> list[ProgressItem]:
    filtered = items
    if filters.get("discipline"):
        filtered = [item for item in filtered if display_text(item.discipline, "未填写专业") == filters["discipline"]]
    if filters.get("building"):
        filtered = [item for item in filtered if display_text(item.building, "未填写楼栋") == filters["building"]]
    if filters.get("floor"):
        filtered = [item for item in filtered if display_text(item.floor, "未填写楼层") == filters["floor"]]
    if filters.get("delay_level"):
        filtered = [item for item in filtered if _delay_level_matches(_delay_level_value(item), filters["delay_level"])]
    return filtered


def _filter_dashboard_scope_items(items: list[ProgressItem], filters: dict[str, str | None]) -> list[ProgressItem]:
    filtered = items
    if filters.get("construction_unit"):
        filtered = [item for item in filtered if display_text(getattr(item, "construction_unit", None), "未填写施工单位") == filters["construction_unit"]]
    if filters.get("discipline"):
        filtered = [item for item in filtered if display_text(item.discipline, "未填写专业") == filters["discipline"]]
    if filters.get("building"):
        filtered = [item for item in filtered if display_text(item.building, "未填写楼栋") == filters["building"]]
    if filters.get("floor"):
        filtered = [item for item in filtered if display_text(item.floor, "未填写楼层") == filters["floor"]]
    if filters.get("system_name"):
        filtered = [item for item in filtered if display_text(item.system_name, "未填写系统") == filters["system_name"]]
    if filters.get("delay_level"):
        filtered = [
            item
            for item in filtered
            if _delay_level_matches(_delay_level_value(item), filters["delay_level"]) or item.status == filters["delay_level"]
        ]
    return filtered


def _filter_dashboard_rectifications(items: list[RectificationItem], filters: dict[str, str | None]) -> list[RectificationItem]:
    filtered = items
    if filters.get("construction_unit"):
        filtered = [item for item in filtered if display_text(item.responsible_unit, "未填写施工单位") == filters["construction_unit"]]
    for field in ("discipline", "building", "floor", "system_name"):
        if filters.get(field):
            filtered = [item for item in filtered if display_text(getattr(item, field), "未填写") == filters[field]]
    if filters.get("delay_level"):
        filtered = [item for item in filtered if _delay_level_matches(item.delay_level, filters["delay_level"])]
    return filtered


def _dashboard_filter_text(filters: dict[str, str | None]) -> str:
    labels = {
        "construction_unit": "施工单位",
        "building": "楼栋",
        "floor": "楼层",
        "discipline": "专业",
        "system_name": "系统",
        "delay_level": "状态",
    }
    parts = [f"{labels.get(key, key)}={value}" for key, value in filters.items() if value]
    return "；".join(parts) if parts else "全部"


def _delay_level_value(item: ProgressItem) -> str:
    if _is_not_started_for_delay(item):
        return "not_started_by_plan"
    level, _ = delay_level_for_deviation(item.progress_deviation)
    return level or item.status or "unknown"


def _delay_level_matches(actual: str | None, requested: str | None) -> bool:
    if actual == requested:
        return True
    if requested == "delayed_or_worse":
        return actual in {"seriously_delayed", "delayed"}
    if requested == "any_delayed":
        return actual in {"seriously_delayed", "delayed", "slightly_delayed"}
    return False


def _clean_filter(value: str | None) -> str | None:
    if value is None:
        return None
    cleaned = str(value).strip()
    return cleaned or None


def _resolve_report_batches(
    db: Session,
    project_id: int,
    scope: str | None,
    batch_id: int | None,
    data_date: date | None,
    import_group_id: str | None,
    batch_ids: str | None,
) -> list[ImportBatch]:
    normalized_scope = (scope or "").strip().lower()
    if normalized_scope != "project":
        batch = get_published_batch(db, project_id, batch_id)
        return [batch] if batch is not None else []

    parsed_batch_ids = _parse_report_batch_ids(batch_ids)
    statement = select(ImportBatch).where(
        ImportBatch.project_id == project_id,
        ImportBatch.is_active.is_(True),
        ImportBatch.status == "published",
    )
    if import_group_id:
        statement = statement.where(ImportBatch.import_group_id == import_group_id)
    elif parsed_batch_ids:
        statement = statement.where(ImportBatch.id.in_(parsed_batch_ids))
    elif data_date is not None:
        statement = statement.where(ImportBatch.data_date == data_date)
    elif batch_id is not None:
        statement = statement.where(ImportBatch.id == batch_id)
    return list(db.scalars(statement.order_by(ImportBatch.data_date.asc().nullslast(), ImportBatch.id.asc())))


def _parse_report_batch_ids(value: str | None) -> list[int]:
    if not value:
        return []
    ids: list[int] = []
    for part in value.split(","):
        cleaned = part.strip()
        if not cleaned:
            continue
        ids.append(int(cleaned))
    return list(dict.fromkeys(ids))


def _scoped_report_items(
    db: Session,
    project: Project,
    batches: list[ImportBatch],
    baseline,
    profile=None,
) -> list[ProgressItem]:
    items: list[ProgressItem] = []
    for batch in batches:
        batch_items = apply_time_based_progress(list_items(db, project.id, batch.id), batch, profile)
        items.extend(_apply_plan_start_status(filter_items_by_baseline(batch_items, baseline), batch))
    return items


def _unit_text(items: list[ProgressItem], unit_mixed: bool) -> str:
    units = item_units(items)
    return _unit_text_from_values(units, unit_mixed)


def _unit_text_from_values(units: list[str], unit_mixed: bool) -> str:
    if not units:
        return "-"
    prefix = "单位混杂：" if unit_mixed else ""
    return f"{prefix}{'、'.join(units)}"


def _unit_mixed_note() -> str:
    return "当前分组包含多种单位，建议使用任务平均完成率、权重完成率或产值完成率。"


def _status_label(value: str | None) -> str:
    labels = {
        "completed": "已完成",
        "ahead": "超前",
        "normal": "正常",
        "slightly_delayed": "轻微滞后",
        "delayed": "明显滞后",
        "delayed_or_worse": "明显及以上滞后",
        "any_delayed": "全部滞后",
        "seriously_delayed": "严重滞后",
        "seriously_delay": "严重滞后",
        "not_started_by_plan": "未到计划开始",
        "missing_plan_dates": "缺少计划日期",
        "invalid_plan_dates": "计划日期异常",
        "unknown": "未知",
    }
    return labels.get(value or "unknown", "未知")


def _display_value(value, empty: str = "-") -> str:
    if value is None:
        return empty
    text = str(value).strip()
    if not text or text.lower() in {"none", "null", "undefined"}:
        return empty
    return text


def _low_quality_message(batch: ImportBatch, threshold: float = 70.0) -> str | None:
    if batch.data_quality_score is None or batch.data_quality_score >= threshold:
        return None
    return (
        f"当前数据质量评分 {batch.data_quality_score:.1f} 分，低于建议阈值 {threshold:.1f} 分。"
        " 建议检查字段映射、日期格式、工程量、完成率等校验问题。"
    )


def _task_warning_record_count(db: Session, project_id: int, batch_id: int | None) -> int:
    rows = db.execute(
        select(WarningRecord, WarningRule)
        .outerjoin(WarningRule, WarningRule.id == WarningRecord.rule_id)
        .where(
            WarningRecord.project_id == project_id,
            WarningRecord.batch_id == batch_id,
            WarningRecord.is_resolved.is_(False),
        )
    ).all()
    return sum(1 for record, rule in rows if not is_data_quality_warning_record(record, rule))


def _rectification_suggestion(status: str | None) -> str:
    label = _status_label(status)
    if label == "严重滞后":
        return "建议优先协调资源投入，明确责任人和完成时间，并在下次例会重点复查。"
    if label == "明显滞后":
        return "建议制定赶工措施，跟踪材料、人员和交叉作业影响。"
    return "建议持续跟踪，避免滞后扩大。"


def _responsible_person(item: ProgressItem) -> str:
    keys = ["责任人", "负责人", "责任工程师", "responsible_person"]
    if not item.extra_fields:
        return ""
    try:
        parsed = json.loads(item.extra_fields)
    except json.JSONDecodeError:
        return ""
    if not isinstance(parsed, dict):
        return ""
    for key in keys:
        value = parsed.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def _rectification_items(db: Session, project_id: int, batch_id: int | None = None, batch_ids: list[int] | None = None) -> list[RectificationItem]:
    statement = select(RectificationItem).where(RectificationItem.project_id == project_id)
    if batch_ids:
        statement = statement.where(RectificationItem.batch_id.in_(batch_ids))
    elif batch_id is not None:
        statement = statement.where((RectificationItem.batch_id == batch_id) | (RectificationItem.batch_id.is_(None)))
    return list(db.execute(statement.order_by(RectificationItem.created_at.desc(), RectificationItem.id.desc())).scalars())


def _rectification_summary(items: list[RectificationItem]) -> dict[str, int]:
    now = datetime.now()
    week_start_date = date.today() - timedelta(days=date.today().weekday())
    week_start = datetime.combine(week_start_date, datetime.min.time())
    return {
        "total": len(items),
        "open": sum(1 for item in items if item.status == "open"),
        "in_progress": sum(1 for item in items if item.status == "in_progress"),
        "completed": sum(1 for item in items if item.status == "completed"),
        "closed": sum(1 for item in items if item.status == "closed"),
        "ignored": sum(1 for item in items if item.status == "ignored"),
        "overdue": sum(1 for item in items if _rectification_overdue(item)),
        "serious": sum(1 for item in items if item.delay_level in {"seriously_delayed", "seriously_delay", "critical"}),
        "new_this_week": sum(1 for item in items if item.created_at and item.created_at >= week_start),
        "closed_this_week": sum(1 for item in items if item.closed_at and item.closed_at >= week_start),
    }


def _rectification_overdue(item: RectificationItem) -> bool:
    return bool(item.planned_finish_date and item.planned_finish_date < date.today() and item.status not in {"closed", "ignored", "completed"})


def _rectification_status_label(value: str | None) -> str:
    return {
        "open": "未开始",
        "in_progress": "整改中",
        "completed": "已完成",
        "closed": "已关闭",
        "ignored": "已忽略",
    }.get(value or "", value or "-")


def _rectification_delay_label(value: str | None) -> str:
    return {
        "slightly_delayed": "轻微滞后",
        "delayed": "明显滞后",
        "seriously_delayed": "严重滞后",
        "seriously_delay": "严重滞后",
        "critical": "严重滞后",
    }.get(value or "", value or "-")


def _rectification_source_label(value: str | None) -> str:
    return {"manual": "手动创建", "progress_item": "滞后项", "warning": "预警记录"}.get(value or "", value or "-")


def _issue_description(code: str | None, level: str | None) -> str:
    descriptions = {
        "INVALID_DATE": "日期格式可能不正确",
        "SUMMARY_ROW_SKIPPED": "汇总行已跳过",
        "negative_quantity": "工程量字段为负数",
        "percent_out_of_range": "百分比超出合理范围",
        "actual_exceeds_total": "实际完成量超过总工程量",
    }
    return descriptions.get(code or "", f"{level or 'unknown'} 校验问题")


def _percent_text(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):.1f}%"


def _weight_percent_text(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{float(value) * 100:.2f}%"


def _signed_percent_text(value: float | None) -> str:
    if value is None:
        return "-"
    number = float(value)
    return f"{'+' if number > 0 else ''}{number:.1f}%"


def _signed_number_text(value: float | None) -> str:
    if value is None:
        return "-"
    number = float(value)
    return f"{'+' if number > 0 else ''}{number:.1f}"


def _safe_file_part(value: str) -> str:
    text = re.sub(r'[\\/:*?"<>|]+', "_", value).strip(" _.")
    text = re.sub(r"_+", "_", text)
    return text or "项目"


def _style_workbook(workbook: Workbook) -> None:
    header_fill = PatternFill("solid", fgColor="DCE4EE")
    for sheet in workbook.worksheets:
        header_row_index = 11 if sheet.title == "整改清单" else 2
        sheet.freeze_panes = f"A{header_row_index + 1}"
        for row_index in (1, header_row_index):
            for cell in sheet[row_index]:
                cell.font = Font(bold=True)
                cell.fill = header_fill
                cell.alignment = Alignment(vertical="center", wrap_text=True)
        if sheet.title == "整改清单":
            sheet["A1"].font = Font(bold=True, size=14)
            for row_index in range(2, 10):
                sheet[f"A{row_index}"].font = Font(bold=True)
        for column_cells in sheet.columns:
            max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
            width_limit = 80 if sheet.title == "进度分析说明" else (60 if sheet.title in {"滞后项清单", "整改清单", "数据质量与校验问题汇总"} else 36)
            column_letter = column_cells[0].column_letter
            width = min(max(max_length + 2, 12), width_limit)
            if sheet.title == "进度分析说明":
                width = 18 if column_letter == "A" else 80
            if sheet.title == "整改清单" and column_letter in {"K", "L"}:
                width = 46
            sheet.column_dimensions[column_letter].width = width
            for cell in column_cells:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
